"""
量子知识图谱推演引擎 - FastAPI 后端服务
==================================================

支持两种核心模式：
1. 推理推演 (reasoning): 基于子图的 Chain-of-Thought 逐步推理，共5步
2. 多智能体模拟 (simulation): 多智能体社交推演，共5个阶段

仿照 MiroFish 流程但适配量子计算领域。
"""

import json
import asyncio
import re
import os
import time
import tempfile
from contextvars import ContextVar
from typing import List, Dict, Any, Tuple, Optional
from collections import defaultdict, Counter
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, HTMLResponse, FileResponse
from pydantic import BaseModel, Field
import httpx

# ============================================================
# 配置
# ============================================================

DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY", "")
DEEPSEEK_BASE_URL = "https://api.deepseek.com"
DEEPSEEK_MODEL = "deepseek-v4-pro"
DEFAULT_HOST = os.environ.get("HOST", "0.0.0.0")
DEFAULT_PORT = int(os.environ.get("PORT", "6122"))
ACTIVE_MODEL_CONFIG: ContextVar[Optional[Dict[str, Any]]] = ContextVar("ACTIVE_MODEL_CONFIG", default=None)

MODEL_PROVIDER_PRESETS = [
    {
        "id": "deepseek", "label": "DeepSeek", "api_kind": "openai",
        "base_url": "https://api.deepseek.com", "models": ["deepseek-v4-pro", "deepseek-v4-flash", "deepseek-chat", "deepseek-reasoner"]
    },
    {
        "id": "openai", "label": "OpenAI", "api_kind": "openai_responses",
        "base_url": "https://api.openai.com/v1", "models": ["gpt-5.5", "gpt-5.4", "gpt-5.4-mini", "gpt-5.4-nano", "gpt-4.1"]
    },
    {
        "id": "anthropic", "label": "Anthropic Claude", "api_kind": "anthropic",
        "base_url": "https://api.anthropic.com/v1", "models": ["claude-opus-4-8", "claude-sonnet-4-6", "claude-haiku-4-5-20251001", "claude-haiku-4-5"]
    },
    {
        "id": "gemini", "label": "Google Gemini", "api_kind": "gemini",
        "base_url": "https://generativelanguage.googleapis.com/v1beta", "models": ["gemini-3.1-pro", "gemini-3.5-flash", "gemini-3-flash", "gemini-2.5-pro", "gemini-2.5-flash"]
    },
    {
        "id": "qwen", "label": "阿里通义千问 Qwen", "api_kind": "openai",
        "base_url": "https://dashscope.aliyuncs.com/compatible-mode/v1", "models": ["qwen3.5-plus", "qwen3-max", "qwen-plus", "qwen-flash", "qwen3-235b-a22b"]
    },
    {
        "id": "kimi", "label": "Moonshot Kimi", "api_kind": "openai",
        "base_url": "https://api.moonshot.cn/v1", "models": ["kimi-k2.6", "kimi-k2.6-code", "kimi-k2.5", "moonshot-v1-128k"]
    },
    {
        "id": "zhipu", "label": "智谱 GLM", "api_kind": "openai",
        "base_url": "https://open.bigmodel.cn/api/paas/v4", "models": ["glm-4.7", "glm-4.6", "glm-4.5", "glm-4-plus", "glm-4-flash"]
    },
    {
        "id": "doubao", "label": "火山方舟 Doubao", "api_kind": "openai",
        "base_url": "https://ark.cn-beijing.volces.com/api/v3", "models": ["doubao-seed-2-0-pro-260215", "doubao-seed-1-8", "doubao-seed-1-6", "doubao-seed-code"]
    },
    {
        "id": "siliconflow", "label": "SiliconFlow", "api_kind": "openai",
        "base_url": "https://api.siliconflow.cn/v1", "models": ["deepseek-ai/DeepSeek-V3.2", "Pro/deepseek-ai/DeepSeek-V3.2", "deepseek-ai/DeepSeek-V3.2-Speciale", "Qwen/Qwen3.5-Plus", "THUDM/GLM-4.6"]
    },
    {
        "id": "openrouter", "label": "OpenRouter", "api_kind": "openai",
        "base_url": "https://openrouter.ai/api/v1", "models": ["openrouter/free", "openai/gpt-5.5", "google/gemini-3.5-flash", "anthropic/claude-sonnet-4-6", "deepseek/deepseek-v3.2"]
    },
    {
        "id": "custom", "label": "自定义 OpenAI 兼容接口", "api_kind": "openai",
        "base_url": "", "models": []
    },
]
MODEL_PROVIDER_BY_ID = {p["id"]: p for p in MODEL_PROVIDER_PRESETS}

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
GRAPH_DATA_PATH = os.path.join(BASE_DIR, "graph_data.json")
HTML_FILE_PATH = os.path.join(BASE_DIR, "quantum_reasoning.html")
DEFAULT_EXPORT_DIR = os.path.join(os.path.expanduser("~"), "Downloads")
STATIC_MIME_TYPES = {
    ".html": "text/html; charset=utf-8",
    ".js": "text/javascript; charset=utf-8",
    ".mjs": "text/javascript; charset=utf-8",
    ".css": "text/css; charset=utf-8",
    ".json": "application/json; charset=utf-8",
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".svg": "image/svg+xml",
    ".ico": "image/x-icon",
    ".txt": "text/plain; charset=utf-8",
    ".md": "text/markdown; charset=utf-8",
    ".mp4": "video/mp4",
    ".srt": "text/plain; charset=utf-8",
}

# ============================================================
# 图谱数据加载
# ============================================================

print(f"[图谱] 正在加载: {GRAPH_DATA_PATH}")

try:
    with open(GRAPH_DATA_PATH, "r", encoding="utf-8") as f:
        GRAPH_DATA = json.load(f)
    print(f"[图谱] 加载成功: {len(GRAPH_DATA['nodes'])} 节点, {len(GRAPH_DATA['edges'])} 边")
except Exception as e:
    print(f"[图谱] 加载失败: {e}")
    GRAPH_DATA = {"nodes": [], "edges": []}

NODES: Dict[str, Dict] = {n["id"]: n for n in GRAPH_DATA.get("nodes", [])}
EDGES: List[Dict] = GRAPH_DATA.get("edges", [])

# 邻接表: node_id -> [(edge_dict, neighbor_id)]
ADJ: Dict[str, List[Tuple[Dict, str]]] = defaultdict(list)
for e in EDGES:
    src, tgt = e.get("source"), e.get("target")
    if src and tgt:
        ADJ[src].append((e, tgt))
        ADJ[tgt].append((e, src))

# 标签索引
LABEL_INDEX: Dict[str, str] = {}
for n in GRAPH_DATA.get("nodes", []):
    lb = n["label"].lower()
    LABEL_INDEX[lb] = n["id"]
    base = re.sub(r'[（(].+?[）)]', '', n["label"]).strip().lower()
    if base and base not in LABEL_INDEX:
        LABEL_INDEX[base] = n["id"]

# 类型映射
TYPE_LABELS = {
    "concept": "概念", "algorithm": "算法", "hardware": "硬件",
    "company": "公司", "application": "应用", "person": "人物",
    "protocol": "协议", "tool": "工具", "material": "材料",
    "theory": "理论", "organization": "机构", "event": "事件",
    "product": "产品", "standard": "标准"
}

# ============================================================
# 关键词匹配映射
# ============================================================

KEYWORD_MAP = {
    "超导": ["超导", "superconducting", "transmon", "超导量子比特", "超导电路"],
    "离子阱": ["离子阱", "离子", "ion trap", "iontrap", "paul trap", "penning"],
    "光子": ["光子", "光量子", "photonic", "光学量子", "线性光学"],
    "拓扑": ["拓扑", "topological", "拓扑量子计算", "majorana"],
    "量子比特": ["量子比特", "qubit", "量子位", "物理量子比特", "逻辑量子比特"],
    "纠错": ["纠错", "容错", "QEC", "fault tolerance", "量子纠错", "fault-tolerant"],
    "退相干": ["退相干", "相干", "decoherence", "coherence time", "T1", "T2"],
    "算法": ["算法", "algorithm", "vqe", "qaoa", "shor", "grover", "hhl", "qaia"],
    "模拟": ["模拟", "simulation", "emu", "quantum simulation", "费米子模拟", "化学模拟"],
    "优化": ["优化", "optimization", "组合优化", "量子近似优化"],
    "机器学习": ["机器学习", "ML", "QML", "量子机器学习", "神经网络", "变分量子算法"],
    "密码": ["密码", "加密", "cryptography", "后量子密码", "PQC", "RSA"],
    "芯片": ["芯片", "chip", "处理器", "soc", "量子处理器", "QPU", "量子芯片"],
    "云": ["云", "cloud", "量子云", "云平台", "公有云"],
    "中国": ["中国", "合肥", "北京", "上海", "深圳", "本源", "国仪", "中微达信", "国盾", "百度", "阿里", "腾讯", "华为", "中国科大", "中科大", "清华", "北大"],
    "美国": ["美国", "IBM", "Google", "Microsoft", "Rigetti", "IonQ", "Quantinuum", "Amazon", "AWS", "Braket"],
    "投资": ["投资", "融资", "市场", "market", "invest", "估值", "IPO", "风投", "VC"],
    "趋势": ["趋势", "发展", "未来", "trend", "roadmap", "路线图", "展望"],
    "公司": ["公司", "企业", "company", "startup", "创业公司"],
    "人物": ["人物", "教授", "博士", "专家", "院士", "PI", "科学家"],
    "产业": ["产业", "industrial", "商业化", "commercial", "产业化"],
    "竞争": ["竞争", "competition", "格局", "landscape", "领跑", "领先"],
    "表面码": ["表面码", "surface code", "planar code"],
    "LDPC": ["LDPC", "ldpc", "低密度奇偶校验", "quantum ldpc"],
    "中性原子": ["中性原子", "neutral atom", "optical tweezer", "光镊"],
    "里德伯": ["里德伯", "rydberg", "里德堡"],
    "谷歌": ["谷歌", "Google", "Alphabet", "Willow", "Sycamore"],
    "IBM": ["IBM", "IBM Quantum", "Condor", "Heron", "Osprey"],
    "微软": ["微软", "Microsoft", "Azure Quantum", "拓扑量子比特"],
    "量子体积": ["量子体积", "Quantum Volume", "QV"],
    "门保真度": ["门保真度", "fidelity", "保真度", "gate fidelity", "CLOPS"],
    "量子优势": ["量子优势", "quantum advantage", "quantum supremacy", "量子霸权"],
    "NISQ": ["NISQ", "noisy intermediate-scale", "含噪声中等规模"],
    "变分": ["变分", "variational", "VQE", "QAOA"],
    "量子通信": ["量子通信", "quantum communication", "量子密钥", "QKD"],
    "量子传感": ["量子传感", "quantum sensing", "量子精密测量", "magnetometer"],
    "稀释制冷机": ["稀释制冷机", "dilution refrigerator", "低温", "mK"],
    "控制系统": ["控制系统", "control", "AWG", "任意波形发生器", "FPGA"],
    "软件栈": ["软件栈", "software stack", "Qiskit", "Cirq", "PennyLane", "TKET"],
    "量子纠错码": ["纠错码", "stabilizer", "stabilizer code", "CSS码", "color code"],
    "逻辑门": ["逻辑门", "logical gate", "physical gate", "transversal"],
    "测量": ["测量", "measurement", "readout", "读出"],
    "耦合": ["耦合", "coupling", "电容器耦合", "微波耦合"],
    "封装": ["封装", "packaging", "互连", "interconnect"],
    "日本": ["日本", "Japan", "Fujitsu", "Hitachi", "NTT", "Riken"],
    "欧洲": ["欧洲", "Europe", "EU", "QuantERA", "EuroHPC", "荷兰", "德国", "法国"],
    "人才": ["人才", "talent", "教育", "培养", "博士点"],
    "政策": ["政策", "policy", "国家战略", "量子计划", "量子倡议"],
    "标准": ["标准", "standard", "IEEE", "ISO", "互操作性"],
    "供应链": ["供应链", "supply chain", "设备", "instrument"],
    "应用案例": ["应用案例", "usecase", "金融", "制药", "材料发现"],
}

# ============================================================
# 子图检索
# ============================================================

def find_seed_nodes(question: str, top_k: int = 15) -> List[Dict]:
    """从问题中提取关键词，匹配图谱节点，按重要性排序返回种子节点。"""
    q_lower = question.lower()
    scored = []

    for n in GRAPH_DATA.get("nodes", []):
        score = 0.0
        label = n["label"]
        label_lower = label.lower()
        desc = (n.get("description") or "").lower()
        importance = n.get("properties", {}).get("importance", 3)

        # 1. 精确匹配节点标签
        if label_lower in q_lower:
            score += 20.0
        else:
            q_words = set(re.findall(r'\w+', q_lower))
            label_words = set(re.findall(r'\w+', label_lower))
            if q_words & label_words:
                score += 10.0
            # 2. 标签部分匹配
            elif any(part in q_lower for part in label_lower.split() if len(part) > 2):
                score += 8.0

        # 3. 描述匹配
        for word in re.findall(r'\w+', q_lower):
            if len(word) > 2 and word in desc:
                score += 1.0

        # 4. 关键词映射匹配
        for key, terms in KEYWORD_MAP.items():
            q_has = any(t.lower() in q_lower for t in terms)
            if q_has:
                n_has = any(t.lower() in label_lower or t.lower() in desc for t in terms)
                if n_has:
                    score += 5.0

        # 5. 重要度加权
        score += importance * 0.5

        if score > 0:
            scored.append((n, score))

    scored.sort(key=lambda x: -x[1])
    return [n for n, s in scored[:top_k]]


def get_subgraph(seed_ids: List[str], max_hops: int = 3, max_nodes: int = 80) -> Dict[str, Any]:
    """BFS 扩展获取相关子图。"""
    visited_nodes = set(seed_ids)
    visited_edges = set()
    frontier = set(seed_ids)
    result_nodes = []
    result_edges = []

    for hop in range(max_hops):
        if len(visited_nodes) >= max_nodes:
            break
        next_frontier = set()
        for nid in list(frontier):
            if len(visited_nodes) >= max_nodes:
                break
            for edge, neighbor in ADJ.get(nid, []):
                if len(visited_nodes) >= max_nodes:
                    break
                edge_key = (edge["source"], edge["target"], edge.get("label", ""))
                if edge_key not in visited_edges:
                    visited_edges.add(edge_key)
                    result_edges.append(edge)
                if neighbor not in visited_nodes:
                    visited_nodes.add(neighbor)
                    next_frontier.add(neighbor)
        frontier = next_frontier
        if not frontier:
            break

    for nid in visited_nodes:
        if nid in NODES:
            result_nodes.append(NODES[nid])

    return {"nodes": result_nodes, "edges": result_edges}


def subgraph_to_context(subgraph: Dict[str, Any], lang: str = "zh") -> str:
    """将子图转为 LLM 可读的结构化文本上下文。"""
    is_en = lang == "en"
    lines = ["【Quantum Knowledge Graph Data】" if is_en else "【量子知识图谱相关数据】"]

    lines.append("=== Entity Nodes ===" if is_en else "=== 实体节点 ===")
    for n in subgraph["nodes"]:
        type_label = TYPE_LABELS.get(n["type"], n["type"])
        importance = n.get("properties", {}).get("importance", 3)
        desc = n.get("description", "")
        if is_en:
            lines.append(f"[{type_label}] {n['label']} (Importance:{importance}/5, ID:{n['id']})")
            if desc:
                lines.append(f"  Description: {desc}")
        else:
            lines.append(f"[{type_label}] {n['label']} (重要度:{importance}/5, ID:{n['id']})")
            if desc:
                lines.append(f"  描述: {desc}")

    lines.append("")
    lines.append("=== Relationship Edges ===" if is_en else "=== 关系边 ===")
    node_map = {n["id"]: n for n in subgraph["nodes"]}
    for e in subgraph["edges"]:
        src = node_map.get(e["source"], {}).get("label", e["source"])
        tgt = node_map.get(e["target"], {}).get("label", e["target"])
        label = e.get("label", "related" if is_en else "相关")
        weight = e.get("weight", 1.0)
        if is_en:
            lines.append(f"  {src} --[{label}, intensity:{weight}]--> {tgt}")
        else:
            lines.append(f"  {src} --[{label}, 强度:{weight}]--> {tgt}")

    return "\n".join(lines)


def get_node_neighborhood(node_id: str, max_hops: int = 2) -> Dict[str, Any]:
    """获取某个节点的邻域子图。"""
    return get_subgraph([node_id], max_hops=max_hops, max_nodes=30)


async def safe_sse_stream(generator, context: str):
    """Prevent generator exceptions from becoming opaque browser network errors."""
    try:
        async for chunk in generator:
            yield chunk
    except asyncio.CancelledError:
        raise
    except Exception as e:
        print(f"[SSE:{context}] 未处理异常: {e}")
        msg = json.dumps({
            "type": "error",
            "message": "后端推演流中断",
            "detail": str(e)
        }, ensure_ascii=False)
        yield f"data: {msg}\n\n"


# ============================================================
# 大模型 API 调用
# ============================================================

def normalize_model_config(config: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    """合并用户模型配置与默认 DeepSeek 配置。"""
    config = config or ACTIVE_MODEL_CONFIG.get() or {}
    provider = (config.get("provider") or "deepseek").strip()
    preset = MODEL_PROVIDER_BY_ID.get(provider, MODEL_PROVIDER_BY_ID["deepseek"])
    api_kind = (config.get("api_kind") or preset.get("api_kind") or "openai").strip()
    model = (config.get("model") or (preset.get("models") or [DEEPSEEK_MODEL])[0] or DEEPSEEK_MODEL).strip()
    base_url = (config.get("base_url") or preset.get("base_url") or DEEPSEEK_BASE_URL).strip().rstrip("/")
    api_key = (config.get("api_key") or (DEEPSEEK_API_KEY if provider == "deepseek" else "")).strip()

    return {
        "provider": provider,
        "provider_label": preset.get("label", provider),
        "api_kind": api_kind,
        "model": model,
        "base_url": base_url,
        "api_key": api_key,
    }


def require_model_config(config: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    resolved = normalize_model_config(config)
    missing = []
    if not resolved["api_key"]:
        missing.append("api_key")
    if not resolved["model"]:
        missing.append("model")
    if resolved["api_kind"] != "gemini" and not resolved["base_url"]:
        missing.append("base_url")
    if missing:
        raise HTTPException(status_code=400, detail=f"模型配置不完整，缺少: {', '.join(missing)}")
    return resolved


def openai_chat_url(base_url: str) -> str:
    base = (base_url or "").rstrip("/")
    if base.endswith("/chat/completions"):
        return base
    return f"{base}/chat/completions"


def openai_responses_url(base_url: str) -> str:
    base = (base_url or "").rstrip("/")
    if base.endswith("/responses"):
        return base
    return f"{base}/responses"


def model_timeout() -> httpx.Timeout:
    return httpx.Timeout(connect=30.0, read=120.0, write=30.0, pool=30.0)


def extract_api_error_text(response: httpx.Response) -> str:
    try:
        data = response.json()
        if isinstance(data, dict):
            error = data.get("error") or data.get("detail") or data
            if isinstance(error, dict):
                return str(error.get("message") or error.get("msg") or error.get("type") or error)[:800]
            return str(error)[:800]
    except Exception:
        pass
    return response.text[:800]


def split_system_messages(messages: List[Dict[str, str]]) -> Tuple[str, List[Dict[str, str]]]:
    system_parts = []
    regular_messages = []
    for msg in messages:
        role = msg.get("role", "user")
        content = msg.get("content", "")
        if role == "system":
            system_parts.append(content)
        else:
            regular_messages.append({"role": role if role in {"user", "assistant"} else "user", "content": content})
    return "\n\n".join(system_parts).strip(), regular_messages


async def call_openai_compatible(config: Dict[str, Any], messages, temperature=0.7, max_tokens=4096, stream=False):
    headers = {
        "Authorization": f"Bearer {config['api_key']}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": config["model"],
        "messages": messages,
        "temperature": temperature,
        "max_tokens": max_tokens,
        "stream": stream
    }

    client = httpx.AsyncClient(timeout=model_timeout(), trust_env=False)
    return client, "POST", openai_chat_url(config["base_url"]), headers, payload


def response_text_from_openai_responses(data: Dict[str, Any]) -> str:
    if data.get("output_text"):
        return str(data["output_text"]).strip()
    parts = []
    for item in data.get("output", []) or []:
        for content in item.get("content", []) or []:
            if isinstance(content, dict):
                if content.get("text"):
                    parts.append(str(content.get("text")))
                elif content.get("type") in {"output_text", "text"} and content.get("value"):
                    parts.append(str(content.get("value")))
    return "".join(parts).strip()


def response_text_from_chat_completion(data: Dict[str, Any]) -> str:
    choices = data.get("choices") or []
    if not choices:
        return ""
    message = choices[0].get("message") or {}
    return (message.get("content") or message.get("reasoning_content") or "").strip()


async def call_openai_responses(config: Dict[str, Any], messages, temperature=0.7, max_tokens=4096) -> str:
    headers = {
        "Authorization": f"Bearer {config['api_key']}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": config["model"],
        "input": messages,
        "temperature": temperature,
        "max_output_tokens": max_tokens,
    }
    async with httpx.AsyncClient(timeout=model_timeout(), trust_env=False) as client:
        response = await client.post(openai_responses_url(config["base_url"]), headers=headers, json=payload)
        if response.status_code != 200:
            raise HTTPException(status_code=response.status_code, detail=f"{config['provider_label']} API错误: {extract_api_error_text(response)}")
        data = response.json()
        return response_text_from_openai_responses(data)


async def call_anthropic(config: Dict[str, Any], messages, temperature=0.7, max_tokens=4096) -> str:
    system, regular_messages = split_system_messages(messages)
    payload = {
        "model": config["model"],
        "messages": regular_messages or [{"role": "user", "content": "Hello"}],
        "temperature": temperature,
        "max_tokens": max_tokens,
    }
    if system:
        payload["system"] = system
    headers = {
        "x-api-key": config["api_key"],
        "anthropic-version": "2023-06-01",
        "Content-Type": "application/json"
    }
    async with httpx.AsyncClient(timeout=model_timeout(), trust_env=False) as client:
        response = await client.post(f"{config['base_url'].rstrip('/')}/messages", headers=headers, json=payload)
        if response.status_code != 200:
            raise HTTPException(status_code=response.status_code, detail=f"{config['provider_label']} API错误: {extract_api_error_text(response)}")
        data = response.json()
        return "".join(part.get("text", "") for part in data.get("content", []) if part.get("type") == "text").strip()


async def call_gemini(config: Dict[str, Any], messages, temperature=0.7, max_tokens=4096) -> str:
    system, regular_messages = split_system_messages(messages)
    contents = []
    if system:
        contents.append({"role": "user", "parts": [{"text": f"System instructions:\n{system}"}]})
    for msg in regular_messages:
        role = "model" if msg.get("role") == "assistant" else "user"
        contents.append({"role": role, "parts": [{"text": msg.get("content", "")}]})
    payload = {
        "contents": contents or [{"role": "user", "parts": [{"text": "Hello"}]}],
        "generationConfig": {
            "temperature": temperature,
            "maxOutputTokens": max_tokens
        }
    }
    url = f"{config['base_url'].rstrip('/')}/models/{config['model']}:generateContent?key={config['api_key']}"
    async with httpx.AsyncClient(timeout=model_timeout(), trust_env=False) as client:
        response = await client.post(url, json=payload)
        if response.status_code != 200:
            raise HTTPException(status_code=response.status_code, detail=f"{config['provider_label']} API错误: {extract_api_error_text(response)}")
        data = response.json()
        parts = data.get("candidates", [{}])[0].get("content", {}).get("parts", [])
        return "".join(part.get("text", "") for part in parts).strip()


async def call_deepseek_stream(messages, temperature=0.7, max_tokens=4096):
    """流式调用当前配置的大模型，yield SSE 格式字符串。"""
    try:
        config = require_model_config()
    except HTTPException as e:
        msg = json.dumps({"type": "error", "message": "模型配置错误", "detail": e.detail}, ensure_ascii=False)
        yield f"data: {msg}\n\n"
        return

    if config["api_kind"] in {"anthropic", "gemini", "openai_responses"}:
        try:
            if config["api_kind"] == "anthropic":
                content = await call_anthropic(config, messages, temperature, max_tokens)
            elif config["api_kind"] == "openai_responses":
                content = await call_openai_responses(config, messages, temperature, max_tokens)
            else:
                content = await call_gemini(config, messages, temperature, max_tokens)
            msg = json.dumps({"type": "content", "content": content}, ensure_ascii=False)
            yield f"data: {msg}\n\n"
        except HTTPException as e:
            msg = json.dumps({"type": "error", "message": f"API错误: {e.status_code}", "detail": e.detail}, ensure_ascii=False)
            yield f"data: {msg}\n\n"
        except Exception as e:
            msg = json.dumps({"type": "error", "message": "模型调用异常", "detail": str(e)}, ensure_ascii=False)
            yield f"data: {msg}\n\n"
        return

    try:
        client, method, url, headers, payload = await call_openai_compatible(config, messages, temperature, max_tokens, stream=True)
        async with client:
            async with client.stream(
                method,
                url,
                headers=headers,
                json=payload
            ) as response:
                if response.status_code != 200:
                    error_text = await response.aread()
                    error_detail = error_text.decode()[:500]
                    msg = json.dumps({"type": "error", "message": f"API错误: {response.status_code}", "detail": error_detail}, ensure_ascii=False)
                    yield f"data: {msg}\n\n"
                    return

                async for line in response.aiter_lines():
                    if not line.startswith("data: "):
                        continue
                    data_str = line[6:]
                    if data_str.strip() == "[DONE]":
                        return
                    try:
                        data = json.loads(data_str)
                        choices = data.get("choices", [{}])
                        if not choices:
                            continue
                        delta = choices[0].get("delta", {})
                        content = delta.get("content", "")
                        if content:
                            msg = json.dumps({"type": "content", "content": content}, ensure_ascii=False)
                            yield f"data: {msg}\n\n"
                    except json.JSONDecodeError:
                        continue
    except httpx.TimeoutException as e:
        msg = json.dumps({"type": "error", "message": "模型响应超时", "detail": str(e)}, ensure_ascii=False)
        yield f"data: {msg}\n\n"
    except httpx.HTTPError as e:
        msg = json.dumps({"type": "error", "message": "模型连接异常", "detail": str(e)}, ensure_ascii=False)
        yield f"data: {msg}\n\n"
    except Exception as e:
        msg = json.dumps({"type": "error", "message": "模型调用异常", "detail": str(e)}, ensure_ascii=False)
        yield f"data: {msg}\n\n"


async def call_deepseek(messages, temperature=0.7, max_tokens=4096) -> str:
    """非流式调用当前配置的大模型，返回完整响应文本。"""
    config = require_model_config()
    if config["api_kind"] == "anthropic":
        return await call_anthropic(config, messages, temperature, max_tokens)
    if config["api_kind"] == "gemini":
        return await call_gemini(config, messages, temperature, max_tokens)
    if config["api_kind"] == "openai_responses":
        return await call_openai_responses(config, messages, temperature, max_tokens)

    client, method, url, headers, payload = await call_openai_compatible(config, messages, temperature, max_tokens, stream=False)
    async with client:
        response = await client.request(method, url, headers=headers, json=payload)
        if response.status_code != 200:
            error_text = extract_api_error_text(response)
            raise HTTPException(status_code=response.status_code, detail=f"{config['provider_label']} API错误: {error_text}")
        data = response.json()
        return response_text_from_chat_completion(data)


async def stream_with_model_config(generator, model_config: Optional[Dict[str, Any]], context: str):
    """为一次 SSE 请求绑定用户选择的大模型配置。"""
    token = ACTIVE_MODEL_CONFIG.set(model_config or None)
    try:
        async for chunk in safe_sse_stream(generator, context):
            yield chunk
    finally:
        ACTIVE_MODEL_CONFIG.reset(token)


# ============================================================
# 推理推演模式（Reasoning Mode）
# ============================================================
#
# SSE 事件格式：
#   step_start: {"type":"step_start","step":1,"title":"..."}
#   step_content: {"type":"step_content","step":1,"content":"..."}
#   step_end: {"type":"step_end","step":1}
#   complete: {"type":"complete","conclusion":"...","evidence":[...]}
#

REASONING_SYSTEM_PROMPT_ZH = """你是量子计算领域的资深专家分析师，精通量子计算的技术、产业、学术和应用各个维度。

你拥有一个量子计算知识图谱的数据，你需要基于图谱数据进行严格的逻辑推理和预测分析。

## 推理规则

你必须按以下步骤进行推理，每一步都要明确标注：

**第1步：问题解析**
- 分析用户问题的核心意图和关注点
- 识别涉及的关键领域（技术/产业/学术/应用/政策）
- 明确需要回答的核心问题

**第2步：知识检索与梳理**
- 基于提供的知识图谱子图，系统列出所有相关实体和关系
- 标注每个实体的类型、重要度和关键属性
- 梳理实体间的关系链，构建知识网络

**第3步：Chain-of-Thought 逻辑推演**
- 从已知事实出发，沿着关系链逐步推理
- 每个子推论都要有图谱数据支撑，引用具体实体ID和关系
- 标注推理链条：实体A →[关系]→ 实体B →[关系]→ 实体C
- 识别关键节点的作用和影响力

**第4步：多维度综合分析**
- 技术维度：技术成熟度、技术路线对比、技术瓶颈
- 产业维度：竞争格局、市场趋势、商业化路径
- 学术维度：研究前沿、突破性进展、未来方向
- 政策维度：政策支持、国家战略、国际竞争

**第5步：预测结论**
- 给出明确、可操作的预测结论
- 标注置信度（高/中/低，并给出数值0-1）
- 列出支撑证据（引用图谱中的具体实体和关系，格式：实体名[ID]）
- 列出不确定因素和风险提示
- 如适用，给出时间线和里程碑预测

## 输出要求
- 使用 Markdown 格式
- 用 ## 标注步骤标题
- 用 > 标注关键推断和结论
- 用 **加粗** 标注关键实体和概念

## 重要约束
- 所有推理必须基于图谱数据，不要编造不存在的关系或事实
- 如果图谱数据不足以支撑某项推断，明确说明"图谱数据不足，此部分为推测"
- 严格区分"事实"和"推测"，推测需标注依据和置信度"""


REASONING_SYSTEM_PROMPT_EN = """You are a senior expert analyst in the field of quantum computing, with deep expertise in technology, industry, academia, and applications.

You have access to a quantum computing knowledge graph dataset. You must perform rigorous logical reasoning and predictive analysis based on this graph data.

## Reasoning Rules

You MUST follow these steps, clearly labeling each one:

**Step 1: Problem Analysis**
- Analyze the core intent and focus of the user's question
- Identify key domains involved (technology/industry/academy/application/policy)
- Clarify the core question that needs to be answered

**Step 2: Knowledge Retrieval & Organization**
- Based on the provided knowledge graph subgraph, systematically list all relevant entities and relationships
- Label each entity's type, importance, and key attributes
- Organize relationship chains between entities to build a knowledge network

**Step 3: Chain-of-Thought Logical Reasoning**
- Start from known facts and reason step by step along relationship chains
- Each sub-inference must be supported by graph data, citing specific entity IDs and relationships
- Label reasoning chains: Entity A →[relation]→ Entity B →[relation]→ Entity C
- Identify the role and influence of key nodes

**Step 4: Multi-dimensional Comprehensive Analysis**
- Technology dimension: maturity level, technical route comparison, bottlenecks
- Industry dimension: competitive landscape, market trends, commercialization paths
- Academic dimension: research frontiers, breakthroughs, future directions
- Policy dimension: policy support, national strategy, international competition

**Step 5: Prediction Conclusion**
- Provide clear, actionable prediction conclusions
- Indicate confidence level (high/medium/low, with numerical value 0-1)
- List supporting evidence (cite specific entities and relationships from the graph, format: EntityName[ID])
- List uncertainties and risk factors
- Where applicable, provide timeline and milestone predictions

## Output Requirements
- Use Markdown format
- Use ## for step titles
- Use > for key inferences and conclusions
- Use **bold** for key entities and concepts

## Important Constraints
- ALL reasoning must be based on graph data. Do not fabricate non-existent relationships or facts
- If graph data is insufficient to support an inference, explicitly state "Insufficient graph data - this is speculative"
- Clearly distinguish "facts" from "speculations", and label speculation with basis and confidence level"""


async def reasoning_mode(question: str, max_nodes: int = 80, lang: str = "zh"):
    """推理推演模式 - 生成 SSE 事件流。"""

    is_en = lang == "en"
    system_prompt = REASONING_SYSTEM_PROMPT_EN if is_en else REASONING_SYSTEM_PROMPT_ZH

    # 第0步标题
    step_titles = {
        0: ("Retrieving Relevant Subgraph", "检索相关子图"),
        1: ("Problem Analysis", "问题解析"),
        2: ("Knowledge Retrieval & Organization", "知识检索与梳理"),
        3: ("Chain-of-Thought Reasoning", "Chain-of-Thought 逻辑推演"),
        4: ("Multi-dimensional Analysis", "多维度综合分析"),
        5: ("Prediction Conclusion", "预测结论"),
    }
    def st(n): return step_titles[n][0] if is_en else step_titles[n][1]

    # 第0步：检索子图
    msg = json.dumps({"type": "step_start", "step": 0, "title": st(0)}, ensure_ascii=False)
    yield f"data: {msg}\n\n"

    analyzing_text = f"Analyzing question: {question}" if is_en else f"正在分析问题：{question}"
    msg = json.dumps({"type": "step_content", "step": 0, "content": analyzing_text}, ensure_ascii=False)
    yield f"data: {msg}\n\n"

    seeds = find_seed_nodes(question, top_k=15)
    if not seeds:
        seeds = sorted(GRAPH_DATA.get("nodes", []), key=lambda n: n.get("properties", {}).get("importance", 0), reverse=True)[:8]

    seed_ids = [n["id"] for n in seeds]
    subgraph = get_subgraph(seed_ids, max_nodes=max_nodes)

    content_text = f"Found {len(seeds)} seed nodes, expanded to {len(subgraph['nodes'])} nodes and {len(subgraph['edges'])} edges in subgraph" if is_en else f"找到 {len(seeds)} 个种子节点，扩展得到 {len(subgraph['nodes'])} 个节点和 {len(subgraph['edges'])} 条边的子图"
    msg = json.dumps({"type": "step_content", "step": 0, "content": content_text}, ensure_ascii=False)
    yield f"data: {msg}\n\n"

    # 推演联动：推送高亮节点事件
    highlight_node_ids = [n["id"] for n in seeds[:10]]
    highlight_edge_pairs = []
    for e in subgraph["edges"][:20]:
        if e["source"] in highlight_node_ids and e["target"] in highlight_node_ids:
            highlight_edge_pairs.append([e["source"], e["target"]])
    msg = json.dumps({"type": "highlight_nodes", "node_ids": highlight_node_ids, "edge_pairs": highlight_edge_pairs}, ensure_ascii=False)
    yield f"data: {msg}\n\n"

    msg = json.dumps({"type": "step_end", "step": 0}, ensure_ascii=False)
    yield f"data: {msg}\n\n"

    context = subgraph_to_context(subgraph, lang)

    # 第1步：问题解析
    msg = json.dumps({"type": "step_start", "step": 1, "title": st(1)}, ensure_ascii=False)
    yield f"data: {msg}\n\n"

    if is_en:
        parse_prompt = f"Analyze the following user question and complete Step 1 (Problem Analysis):\n\nUser question: {question}\n\nPlease answer: 1. What is the core intent? 2. Which key domains are involved? 3. What is the core question to answer?\n\nRequirements: Concise and clear, Markdown format."
        messages_parse = [
            {"role": "system", "content": "You are a quantum computing problem analysis expert. Please only complete Step 1 problem analysis."},
            {"role": "user", "content": parse_prompt}
        ]
    else:
        parse_prompt = f"请分析以下用户问题，完成第1步（问题解析）：\n\n用户问题：{question}\n\n请回答：1.问题的核心意图？2.涉及哪些关键领域？3.需要回答的核心问题？\n\n要求：简洁明确，Markdown格式。"
        messages_parse = [
            {"role": "system", "content": "你是量子计算领域的问题分析专家，请只完成第1步的问题解析。"},
            {"role": "user", "content": parse_prompt}
        ]

    step1_content = ""
    async for chunk in call_deepseek_stream(messages_parse, temperature=0.3, max_tokens=1024):
        try:
            parsed = json.loads(chunk.strip().removeprefix("data: "))
            if parsed.get("type") == "content":
                step1_content += parsed["content"]
                yield chunk
            else:
                yield chunk
                if parsed.get("type") == "error":
                    return
        except (json.JSONDecodeError, AttributeError):
            yield chunk

    msg = json.dumps({"type": "step_end", "step": 1}, ensure_ascii=False)
    yield f"data: {msg}\n\n"

    # 第2步：知识检索与梳理
    msg = json.dumps({"type": "step_start", "step": 2, "title": st(2)}, ensure_ascii=False)
    yield f"data: {msg}\n\n"

    type_counts = Counter(n["type"] for n in subgraph["nodes"])

    if is_en:
        node_summary = f"Retrieved {len(subgraph['nodes'])} entity nodes, {len(subgraph['edges'])} relationship edges.\n"
        node_summary += "Entity type distribution:\n"
        for t, cnt in type_counts.most_common():
            node_summary += f"- {TYPE_LABELS.get(t, t)}: {cnt}\n"
    else:
        node_summary = f"共检索到 {len(subgraph['nodes'])} 个实体节点，{len(subgraph['edges'])} 条关系边。\n"
        node_summary += "实体类型分布：\n"
        for t, cnt in type_counts.most_common():
            node_summary += f"- {TYPE_LABELS.get(t, t)}: {cnt} 个\n"

    msg = json.dumps({"type": "step_content", "step": 2, "content": node_summary}, ensure_ascii=False)
    yield f"data: {msg}\n\n"

    msg = json.dumps({"type": "step_end", "step": 2}, ensure_ascii=False)
    yield f"data: {msg}\n\n"

    # 第3步：Chain-of-Thought 推理
    msg = json.dumps({"type": "step_start", "step": 3, "title": st(3)}, ensure_ascii=False)
    yield f"data: {msg}\n\n"

    if is_en:
        cot_prompt = f"Based on the following quantum knowledge graph data, perform Chain-of-Thought logical reasoning (Step 3):\n\n{context}\n\n---\n\nUser question: {question}\n\nComplete Step 3 reasoning. Cite specific entities and relationships from the graph."
        messages_cot = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": cot_prompt}
        ]
    else:
        cot_prompt = f"请基于以下量子知识图谱数据，对问题进行 Chain-of-Thought 逻辑推演（第3步）：\n\n{context}\n\n---\n\n用户问题：{question}\n\n请完成第3步推演，每步引用图谱具体实体和关系。"
        messages_cot = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": cot_prompt}
        ]

    step3_content = ""
    async for chunk in call_deepseek_stream(messages_cot, temperature=0.7, max_tokens=4096):
        try:
            parsed = json.loads(chunk.strip().removeprefix("data: "))
            if parsed.get("type") == "content":
                step3_content += parsed["content"]
                yield chunk
            else:
                yield chunk
                if parsed.get("type") == "error":
                    return
        except (json.JSONDecodeError, AttributeError):
            yield chunk

    msg = json.dumps({"type": "step_end", "step": 3}, ensure_ascii=False)
    yield f"data: {msg}\n\n"

    # 推演联动：扫描推理内容中引用的节点，推送高亮事件
    referenced_node_ids = []
    for n in subgraph["nodes"][:30]:
        if n["label"] in step3_content:
            referenced_node_ids.append(n["id"])
    if referenced_node_ids:
        # 找出引用节点之间的边
        ref_edge_pairs = []
        ref_set = set(referenced_node_ids)
        for e in subgraph["edges"][:30]:
            if e["source"] in ref_set and e["target"] in ref_set:
                ref_edge_pairs.append([e["source"], e["target"]])
        msg = json.dumps({"type": "highlight_nodes", "node_ids": referenced_node_ids, "edge_pairs": ref_edge_pairs}, ensure_ascii=False)
        yield f"data: {msg}\n\n"

    # 第4步：多维度综合分析
    msg = json.dumps({"type": "step_start", "step": 4, "title": st(4)}, ensure_ascii=False)
    yield f"data: {msg}\n\n"

    if is_en:
        analysis_prompt = f"Based on the previous reasoning (Step 3), now perform multi-dimensional comprehensive analysis (Step 4).\n\nPrevious reasoning content:\n{step3_content[:2000]}\n\nUser question: {question}\n\nPlease complete Step 4 analysis across four dimensions: technology/industry/academy/policy."
        messages_analysis = [
            {"role": "system", "content": "You are a senior quantum computing analyst. Please complete multi-dimensional comprehensive analysis."},
            {"role": "user", "content": analysis_prompt}
        ]
    else:
        analysis_prompt = f"基于前面的推理（第3步），现在进行多维度综合分析（第4步）。\n\n前面的推理内容：\n{step3_content[:2000]}\n\n用户问题：{question}\n\n请完成第4步分析：技术/产业/学术/政策四个维度。"
        messages_analysis = [
            {"role": "system", "content": "你是量子计算领域的资深分析师，请完成多维度综合分析。"},
            {"role": "user", "content": analysis_prompt}
        ]

    step4_content = ""
    async for chunk in call_deepseek_stream(messages_analysis, temperature=0.7, max_tokens=4096):
        try:
            parsed = json.loads(chunk.strip().removeprefix("data: "))
            if parsed.get("type") == "content":
                step4_content += parsed["content"]
                yield chunk
            else:
                yield chunk
                if parsed.get("type") == "error":
                    return
        except (json.JSONDecodeError, AttributeError):
            yield chunk

    msg = json.dumps({"type": "step_end", "step": 4}, ensure_ascii=False)
    yield f"data: {msg}\n\n"

    # 第5步：预测结论
    msg = json.dumps({"type": "step_start", "step": 5, "title": st(5)}, ensure_ascii=False)
    yield f"data: {msg}\n\n"

    if is_en:
        conclusion_prompt = f"Based on all previous analysis, provide the final prediction conclusion (Step 5).\n\nProblem analysis: {step1_content[:500]}\nReasoning: {step3_content[:1000]}\nAnalysis: {step4_content[:1000]}\n\nUser question: {question}\n\nPlease complete Step 5: Conclusion, confidence, evidence, uncertainties."
        messages_conclusion = [
            {"role": "system", "content": "You are a senior quantum computing prediction expert. Please provide rigorous prediction conclusions."},
            {"role": "user", "content": conclusion_prompt}
        ]
    else:
        conclusion_prompt = f"基于前面的所有分析，给出最终预测结论（第5步）。\n\n问题解析：{step1_content[:500]}\n推理：{step3_content[:1000]}\n分析：{step4_content[:1000]}\n\n用户问题：{question}\n\n请完成第5步：结论、置信度、证据、不确定因素。"
        messages_conclusion = [
            {"role": "system", "content": "你是量子计算领域的资深预测专家，请给出严谨的预测结论。"},
            {"role": "user", "content": conclusion_prompt}
        ]

    step5_content = ""
    async for chunk in call_deepseek_stream(messages_conclusion, temperature=0.7, max_tokens=4096):
        try:
            parsed = json.loads(chunk.strip().removeprefix("data: "))
            if parsed.get("type") == "content":
                step5_content += parsed["content"]
                yield chunk
            else:
                yield chunk
                if parsed.get("type") == "error":
                    return
        except (json.JSONDecodeError, AttributeError):
            yield chunk

    msg = json.dumps({"type": "step_end", "step": 5}, ensure_ascii=False)
    yield f"data: {msg}\n\n"

    # 推演联动：扫描预测结论中引用的节点
    conclusion_node_ids = []
    for n in subgraph["nodes"][:30]:
        if n["label"] in step5_content:
            conclusion_node_ids.append(n["id"])
    if conclusion_node_ids:
        con_edge_pairs = []
        con_set = set(conclusion_node_ids)
        for e in subgraph["edges"][:20]:
            if e["source"] in con_set and e["target"] in con_set:
                con_edge_pairs.append([e["source"], e["target"]])
        msg = json.dumps({"type": "highlight_nodes", "node_ids": conclusion_node_ids, "edge_pairs": con_edge_pairs}, ensure_ascii=False)
        yield f"data: {msg}\n\n"

    # 完成：推送完整结果
    evidence_ids = re.findall(r'\[([a-z_0-9]+)\]', step5_content)
    evidence = list(set(evidence_ids))[:10]
    confidence_match = re.search(r'置信度.*?([0-9.]+)', step5_content)
    confidence = float(confidence_match.group(1)) if confidence_match else 0.7

    prediction_data = {
        "type": "prediction",
        "conclusion": step5_content,
        "evidence": evidence,
        "confidence": confidence,
        "subgraph": {
            "node_count": len(subgraph["nodes"]),
            "edge_count": len(subgraph["edges"]),
            "seed_nodes": [{"id": n["id"], "label": n["label"], "type": n["type"]} for n in seeds]
        }
    }
    msg = json.dumps(prediction_data, ensure_ascii=False)
    yield f"data: {msg}\n\n"

    complete_data = {
        "type": "complete",
        "conclusion": step5_content,
        "evidence": evidence,
        "confidence": confidence,
        "subgraph": prediction_data["subgraph"]
    }
    msg = json.dumps(complete_data, ensure_ascii=False)
    yield f"data: {msg}\n\n"


# ============================================================
# 多智能体模拟模式（Simulation Mode）
# ============================================================
#
# SSE 事件格式：
#   phase: {"type":"phase","phase":"entity_extraction","message":"..."}
#   agent_profile: {"type":"agent_profile","agent":{...}}
#   round_start: {"type":"round_start","round":1}
#   agent_speech: {"type":"agent_speech","round":1,"agent_id":0,"agent_name":"IBM","content":"..."}
#   round_summary: {"type":"round_summary","round":1,"summary":"..."}
#   coordinator_summary: {"type":"coordinator_summary","content":"..."}
#   prediction: {"type":"prediction","conclusion":"...","confidence":0.75,"evidence":[...]}
#

SIMULATION_SYSTEM_PROMPT_ZH = """你是多智能体推演系统的协调者和主持人。你需要组织和总结多个量子计算领域角色对同一问题的讨论推演。

## 你的职责
1. 组织和引导讨论，确保每位角色都有发言机会
2. 在每轮讨论后做简要总结
3. 在全部讨论结束后，作为协调者给出综合总结和预测

## 讨论规则
- 每个角色基于自己的身份、专业和立场发言
- 角色之间可以有观点冲突，这是正常的
- 发言要符合角色的身份背景，不要跳出角色

## 输出格式（协调者总结）
## 协调者总结
### 各方观点回顾
### 共识
### 分歧
### 综合预测
- 预测结论
- 置信度：X/1.0
- 关键证据
- 不确定因素"""


SIMULATION_SYSTEM_PROMPT_EN = """You are the coordinator and moderator of a multi-agent reasoning system. You need to organize and summarize discussions from multiple quantum computing domain personas on the same question.

## Your Responsibilities
1. Organize and guide the discussion, ensuring each persona has a chance to speak
2. Provide a brief summary after each round of discussion
3. After all discussion rounds, as coordinator provide a comprehensive summary and prediction

## Discussion Rules
- Each persona speaks based on their identity, expertise, and stance
- Viewpoint conflicts between personas are normal and expected
- Speech must match the persona's background; do not break character

## Output Format (Coordinator Summary)
## Coordinator Summary
### Review of Perspectives
### Consensus
### Disagreements
### Comprehensive Prediction
- Prediction Conclusion
- Confidence: X/1.0
- Key Evidence
- Uncertainties"""


def determine_simulation_agent_count(subgraph, question) -> int:
    """根据问题复杂度和可用子图规模动态决定参与智能体数量。"""
    available = len([n for n in subgraph.get("nodes", []) if n.get("type") != "report"])
    if available <= 0:
        return 0

    q = (question or "").lower()
    complexity_terms = [
        "哪些", "有哪些", "如何", "为什么", "是否", "对比", "制约", "关联", "路径",
        "设备", "供应链", "产业", "格局", "中美", "中国", "美国", "竞争", "风险",
        "未来", "趋势", "roadmap", "supply chain", "competition", "risk"
    ]
    matched_terms = sum(1 for term in complexity_terms if term.lower() in q)
    matched_types = len({n.get("type") for n in subgraph.get("nodes", []) if n.get("type")})

    target = 4
    if len(question or "") >= 24:
        target += 1
    if matched_terms >= 2:
        target += 1
    if matched_terms >= 4:
        target += 1
    if matched_types >= 5:
        target += 1
    if len(subgraph.get("edges", [])) >= 80:
        target += 1
    if any(term in q for term in ["中美", "供应链", "制约", "关联", "设备", "哪些", "有哪些"]):
        target += 1

    return max(3, min(10, available, target))


def _node_question_relevance(node, question) -> float:
    q = (question or "").lower()
    label = node.get("label", "")
    label_lower = label.lower()
    desc = (node.get("description") or "").lower()
    score = 0.0

    if label_lower and label_lower in q:
        score += 20
    if "中美" in q and any(term in label for term in ["中国", "美国"]):
        score += 14
    if ("中国" in q or "china" in q) and "中国" in label:
        score += 10
    if ("美国" in q or "usa" in q or "u.s." in q or "america" in q) and "美国" in label:
        score += 10
    if any(term in q for term in ["设备", "供应链", "制约", "关联"]) and node.get("type") in {"hardware", "material", "tool", "company", "organization"}:
        score += 5

    q_words = {w for w in re.findall(r'[\w\u4e00-\u9fff]+', q) if len(w) > 1}
    label_words = {w for w in re.findall(r'[\w\u4e00-\u9fff]+', label_lower) if len(w) > 1}
    desc_words = {w for w in re.findall(r'[\w\u4e00-\u9fff]+', desc) if len(w) > 1}
    score += len(q_words & label_words) * 6
    score += len(q_words & desc_words) * 2

    for terms in KEYWORD_MAP.values():
        q_has = any(t.lower() in q for t in terms)
        if not q_has:
            continue
        n_has = any(t.lower() in label_lower or t.lower() in desc for t in terms)
        if n_has:
            score += 8

    return score


def extract_entities_for_simulation(subgraph, question, max_agents=None):
    """从子图中动态提取适合作为智能体的实体，避免固定人数和单一类型。"""
    if max_agents is None:
        max_agents = determine_simulation_agent_count(subgraph, question)

    candidates = []
    type_weight = {
        "company": 8, "organization": 7, "person": 7,
        "hardware": 7, "material": 6, "tool": 6,
        "concept": 5, "protocol": 5, "standard": 5,
        "algorithm": 4, "application": 4, "theory": 4,
        "software": 4, "security": 4, "qec": 4,
    }
    unsuitable_types = {"report"}

    nodes_with_score = []
    for n in subgraph.get("nodes", []):
        if n.get("type") in unsuitable_types:
            continue
        importance = n.get("properties", {}).get("importance", 3)
        degree = len(ADJ.get(n.get("id"), []))
        relevance = _node_question_relevance(n, question)
        score = relevance + type_weight.get(n.get("type"), 3) + importance * 2 + min(degree, 18) * 0.4
        nodes_with_score.append((n, score))

    nodes_with_score.sort(key=lambda x: x[1], reverse=True)

    selected_ids = set()
    type_counts = Counter()

    def add_candidate(n):
        neighbors = []
        for edge, neighbor_id in ADJ.get(n["id"], []):
            if neighbor_id in NODES:
                neighbor = NODES[neighbor_id]
                neighbors.append({
                    "label": neighbor["label"],
                    "type": neighbor["type"],
                    "relation": edge.get("label", "相关")
                })
        candidates.append({
            "id": n["id"],
            "label": n["label"],
            "type": n["type"],
            "description": n.get("description", ""),
            "importance": n.get("properties", {}).get("importance", 3),
            "neighbors": neighbors[:10]
        })
        selected_ids.add(n["id"])
        type_counts[n["type"]] += 1

    # 第一遍控制类型密度，让公司、硬件、概念、机构等能共同参与。
    for n, _ in nodes_with_score:
        if len(candidates) >= max_agents:
            break
        if n["id"] in selected_ids:
            continue
        if type_counts[n["type"]] >= 3:
            continue
        add_candidate(n)

    # 第二遍补足数量。
    for n, _ in nodes_with_score:
        if len(candidates) >= max_agents:
            break
        if n["id"] in selected_ids:
            continue
        add_candidate(n)

    return candidates


async def generate_agent_profile(entity, question, lang="zh"):
    """使用 LLM 为单个实体生成智能体人设。"""
    is_en = lang == "en"
    neighbors_text = "\n".join([f"- {n['label']}（{TYPE_LABELS.get(n['type'], n['type'])}）: {n['relation']}" for n in entity.get("neighbors", [])])
    if is_en:
        prompt = f"""Generate a multi-agent persona for the following quantum computing entity for use in discussion simulation.

Entity information:
- Name: {entity['label']}
- Type: {TYPE_LABELS.get(entity['type'], entity['type'])}
- Description: {entity.get('description', 'N/A')}
- Importance: {entity.get('importance', 3)}/5

Related entities in knowledge graph:
{neighbors_text if neighbors_text else '(None)'}

Discussion question: {question}

Generate this entity's agent persona, return as JSON format:
{{
    "name": "Display name used in discussions",
    "role": "Role (e.g., Quantum Hardware CEO, Academic Researcher)",
    "expertise": "Professional field (one sentence)",
    "stance": "optimistic/neutral/pessimistic",
    "persona": "Detailed persona description (within 200 words)"
}}

Return JSON only, no other content."""
    else:
        prompt = f"""请为以下量子计算领域的实体生成一个智能体人设，用于多智能体讨论模拟。

实体信息：
- 名称：{entity['label']}
- 类型：{TYPE_LABELS.get(entity['type'], entity['type'])}
- 描述：{entity.get('description', '无')}
- 重要度：{entity.get('importance', 3)}/5

该实体在知识图谱中的关联实体：
{neighbors_text if neighbors_text else '（无关联实体）'}

讨论问题：{question}

请生成该实体的智能体人设，以 JSON 格式返回：
{{
    "name": "实体名称（用于讨论中的称呼）",
    "role": "角色定位（如：量子硬件公司CEO、学术研究员等）",
    "expertise": "专业领域（一句话描述）",
    "stance": "optimistic/neutral/pessimistic",
    "persona": "详细的人设描述（200字以内）"
}}

只返回 JSON，不要其他内容。"""

    try:
        response_text = await call_deepseek([{"role": "user", "content": prompt}], temperature=0.8, max_tokens=1024)
        json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
        if json_match:
            profile = json.loads(json_match.group())
            return {
                "id": entity["id"],
                "name": profile.get("name", entity["label"]),
                "role": profile.get("role", TYPE_LABELS.get(entity["type"], entity["type"])),
                "expertise": profile.get("expertise", ""),
                "stance": profile.get("stance", "neutral"),
                "persona": profile.get("persona", ""),
                "entity_type": entity["type"],
                "entity_label": entity["label"]
            }
    except Exception as e:
        print(f"[人设生成] 失败 {entity['label']}: {e}")

    return {
        "id": entity["id"],
        "name": entity["label"],
        "role": TYPE_LABELS.get(entity["type"], entity["type"]),
        "expertise": entity.get("description", "")[:100],
        "stance": "neutral",
        "persona": f"我是{entity['label']}的代表，从{TYPE_LABELS.get(entity['type'], entity['type'])}的角度分析问题。",
        "entity_type": entity["type"],
        "entity_label": entity["label"]
    }


async def generate_agent_speech(agent, round_num, question, conversation_history, subgraph_context, lang="zh"):
    """使用 LLM 生成单个智能体在一轮讨论中的发言。"""
    is_en = lang == "en"
    history_text = conversation_history if conversation_history else ("(Discussion just started, you are the first to speak)" if is_en else "（讨论刚开始，你是第一个发言者）")
    if is_en:
        prompt = f"""You are {agent['name']}, {agent['role']}.

Your persona:
{agent['persona']}

Your expertise: {agent['expertise']}
Your stance: {agent['stance']}

Discussion question: {question}

Discussion history so far:
{history_text}

Speak in Round {round_num} of the discussion. Requirements:
1. Speech must match your identity, expertise, and stance
2. Respond to previous personas' speeches (if any)
3. Speech length 200-400 words, use first person
4. Stay in character

Output your speech directly, no prefix."""
    else:
        prompt = f"""你是 {agent['name']}，{agent['role']}。

你的人设：
{agent['persona']}

你的专业领域：{agent['expertise']}
你的立场：{agent['stance']}

讨论问题：{question}

至今为止的讨论记录：
{history_text}

请以上面的人设身份，在第{round_num}轮讨论中发言。要求：
1. 发言要符合你的身份、专业和立场
2. 回应前面其他角色的发言（如果有的话）
3. 发言长度200-400字，用第一人称
4. 不要跳出角色

直接输出发言内容，不要加前缀。"""

    try:
        response = await call_deepseek([{"role": "user", "content": prompt}], temperature=0.8, max_tokens=1024)
        return response.strip()
    except Exception as e:
        print(f"[发言生成] 失败 {agent['name']}: {e}")
        return (f"【{agent['name']} speech generation failed】" if is_en else f"【{agent['name']} 发言生成失败】")


async def generate_round_summary(round_num, question, round_speeches, lang="zh"):
    """生成一轮讨论的摘要。"""
    is_en = lang == "en"
    speeches_text = "\n\n".join([f"**{s['agent_name']}**：{s['content']}" for s in round_speeches])
    if is_en:
        prompt = f"Please provide a brief summary of Round {round_num} discussion.\n\nDiscussion question: {question}\n\nThis round's speeches:\n{speeches_text}\n\nSummarize key points, consensus, and disagreements within 200 words."
    else:
        prompt = f"请对第{round_num}轮讨论做简要总结。\n\n讨论问题：{question}\n\n本轮发言：\n{speeches_text}\n\n请用200字以内总结核心观点、共识和分歧。"

    try:
        summary = await call_deepseek([{"role": "user", "content": prompt}], temperature=0.5, max_tokens=512)
        summary = (summary or "").strip()
        if len(summary) < 20 or "摘要生成失败" in summary or "summary generation failed" in summary.lower():
            return build_round_summary_fallback(round_num, round_speeches, lang)
        return summary
    except Exception as e:
        print(f"[轮次摘要] 失败: {e}")
        return build_round_summary_fallback(round_num, round_speeches, lang)


def build_round_summary_fallback(round_num, round_speeches, lang="zh"):
    """当模型摘要为空或失败时，根据本轮发言生成稳定兜底摘要。"""
    is_en = lang == "en"
    if not round_speeches:
        return f"Round {round_num} had no valid speeches to summarize." if is_en else f"第{round_num}轮没有可用于总结的有效发言。"
    names = [s.get("agent_name", "") for s in round_speeches if s.get("agent_name")]
    snippets = []
    for speech in round_speeches[:4]:
        content = re.sub(r"\s+", " ", speech.get("content", "")).strip()
        if content:
            snippets.append(f"{speech.get('agent_name', 'Agent')}: {content[:90]}")
    if is_en:
        return (
            f"Round {round_num} included {len(round_speeches)} agent speeches. "
            f"Main participants included {', '.join(names[:6])}. "
            f"Key positions: {'; '.join(snippets)}. "
            "The round summary was generated locally because the model did not return a complete round summary."
        )
    return (
        f"第{round_num}轮共有{len(round_speeches)}位智能体发言，主要参与方包括{'、'.join(names[:6])}。"
        f"核心观点可概括为：{'；'.join(snippets)}。"
        "本摘要由本地兜底逻辑生成，用于补齐模型未返回完整轮次摘要的情况。"
    )


def _looks_truncated_report(text: str) -> bool:
    """识别明显未写完的预测报告。"""
    text = (text or "").strip()
    if len(text) < 120:
        return True
    return bool(re.search(r"(不是|，|、|：|:|；|;|和|与|但|而|为|在|the|and|or|but|not)$", text, re.IGNORECASE))


def complete_final_prediction_report(question, conclusion, all_speeches, agents, confidence, lang="zh"):
    """为多智能体最终预测提供确定性补全文本，避免模型输出在最后截断。"""
    is_en = lang == "en"
    base = (conclusion or "").strip()
    has_structure = bool(re.search(r"(置信度|Confidence|证据|Evidence|不确定|Uncertainties)", base, re.IGNORECASE))
    if base and not _looks_truncated_report(base) and has_structure:
        return base

    agent_names = [a.get("name", "") for a in agents if a.get("name")]
    round_groups = defaultdict(list)
    for speech in all_speeches:
        round_groups[speech.get("round", 0)].append(speech)
    round_notes = []
    for rnd in sorted(round_groups)[:3]:
        round_notes.append(build_round_summary_fallback(rnd, round_groups[rnd], lang))

    confidence_pct = int(round(float(confidence or 0.75) * 100))
    if is_en:
        return f"""{base}

### Completed Final Prediction
- Prediction: The final answer should prioritize the companies, routes, or research directions that repeatedly appear across technical feasibility, supply-chain resilience, system integration, automated calibration, and commercial operability. A single device advantage is not enough; the strongest candidate is the one that can turn physics insight into a reliable full-stack system.
- Confidence: {confidence_pct}%
- Key evidence: {' '.join(round_notes)}
- Participating perspectives: {', '.join(agent_names[:10])}
- Uncertainties: export-control intensity, long-duration uptime validation, hidden consumable bottlenecks, funding cycles, customer adoption, and whether alternative architectures can reduce dependence on the most sensitive imported components.
"""

    return f"""{base}

### 补齐后的最终预测
- 预测结论：对问题“{question}”的最终判断，不能只看单个器件或单篇论文的领先，而要看谁能同时打通核心设备、供应链韧性、系统集成、校准自动化和商业可运维性。最具胜率的主体通常是具备底层物理理解、光机电一体化、自研控制软件、关键耗材替代和场景牵引能力的团队。
- 置信度：{confidence_pct}%
- 关键证据：{' '.join(round_notes)}
- 参与视角：{'、'.join(agent_names[:10])}
- 不确定因素：出口管制强度、长时间连续运行验证、特种耗材与镀膜工艺自主化、资本投入节奏、下游客户采购意愿，以及替代架构能否真正降低对高端进口部件的依赖。
"""


async def generate_coordinator_summary(question, all_speeches, rounds_count, lang="zh"):
    """协调者总结所有轮次的讨论。"""
    is_en = lang == "en"
    sim_prompt = SIMULATION_SYSTEM_PROMPT_EN if is_en else SIMULATION_SYSTEM_PROMPT_ZH
    speeches_summary = ""
    for s in all_speeches:
        speeches_summary += f"第{s['round']}轮 - {s['agent_name']}：{s['content'][:200]}...\n\n"

    if is_en:
        prompt = f"""You are the coordinator of a multi-agent discussion. Please provide a comprehensive summary of the following multi-round discussion.

Discussion question: {question}
Number of rounds: {rounds_count}

All speech summaries:
{speeches_summary}

Please output in the following format:
## Coordinator Summary
### Review of Perspectives
### Consensus
### Disagreements
### Comprehensive Prediction
- Prediction Conclusion
- Confidence: X/1.0
- Key Evidence
- Uncertainties

Use Markdown format."""
    else:
        prompt = f"""你是多智能体讨论的协调者，请对以下多轮讨论做综合总结。

讨论问题：{question}
讨论轮数：{rounds_count}

所有发言摘要：
{speeches_summary}

请按格式输出：
## 协调者总结
### 各方观点回顾
### 共识
### 分歧
### 综合预测
- 预测结论
- 置信度：X/1.0
- 关键证据
- 不确定因素

用 Markdown 格式。"""

    try:
        return await call_deepseek(
            [{"role": "system", "content": sim_prompt}, {"role": "user", "content": prompt}],
            temperature=0.7, max_tokens=4096
        )
    except Exception as e:
        print(f"[协调者总结] 失败: {e}")
        return ("Coordinator summary generation failed." if is_en else "协调者总结生成失败。")


async def simulation_mode(question, rounds=3, max_nodes=80, lang="zh"):
    """多智能体模拟模式 - 生成 SSE 事件流。"""
    is_en = lang == "en"

    # 阶段1：提取相关实体
    phase_msg_extract = "Extracting relevant entities from knowledge graph..." if is_en else "正在从知识图谱中提取相关实体..."
    msg = json.dumps({"type": "phase", "phase": "entity_extraction", "message": phase_msg_extract}, ensure_ascii=False)
    yield f"data: {msg}\n\n"

    seeds = find_seed_nodes(question, top_k=15)
    if not seeds:
        seeds = sorted(GRAPH_DATA.get("nodes", []), key=lambda n: n.get("properties", {}).get("importance", 0), reverse=True)[:10]

    seed_ids = [n["id"] for n in seeds]
    subgraph = get_subgraph(seed_ids, max_nodes=max_nodes)
    subgraph_context = subgraph_to_context(subgraph, lang)
    entities = extract_entities_for_simulation(subgraph, question)

    msg = json.dumps({"type": "phase", "phase": "entity_extraction", "message": f"{'Extracted ' + str(len(entities)) + ' relevant entities' if is_en else '已提取 ' + str(len(entities)) + ' 个相关实体'}"}, ensure_ascii=False)
    yield f"data: {msg}\n\n"

    # 推演联动：推送高亮节点
    entity_ids = [e["id"] for e in entities]
    entity_edge_pairs = []
    for e in subgraph["edges"][:15]:
        if e["source"] in entity_ids and e["target"] in entity_ids:
            entity_edge_pairs.append([e["source"], e["target"]])
    msg = json.dumps({"type": "highlight_nodes", "node_ids": entity_ids, "edge_pairs": entity_edge_pairs}, ensure_ascii=False)
    yield f"data: {msg}\n\n"

    # 阶段2：生成智能体人设
    phase_msg_profile = "Generating agent profiles for each entity..." if is_en else "正在为每个实体生成智能体人设..."
    msg = json.dumps({"type": "phase", "phase": "profile_generation", "message": phase_msg_profile}, ensure_ascii=False)
    yield f"data: {msg}\n\n"

    agents = []
    for i, entity in enumerate(entities):
        msg_text = (f"Generating profile for {entity['label']}..." if is_en else f"正在为 {entity['label']} 生成人设...")
        msg = json.dumps({"type": "phase", "phase": "profile_generation", "message": msg_text}, ensure_ascii=False)
        yield f"data: {msg}\n\n"

        profile = await generate_agent_profile(entity, question, lang)
        profile["agent_id"] = i
        profile["entity_id"] = entity.get("id", "")
        profile["entity_type"] = entity.get("type", "concept")
        agents.append(profile)

        msg = json.dumps({"type": "agent_profile", "agent": profile}, ensure_ascii=False)
        yield f"data: {msg}\n\n"

    msg = json.dumps({"type": "phase", "phase": "profile_generation", "message": f"{'Generated ' + str(len(agents)) + ' agent profiles' if is_en else '已生成 ' + str(len(agents)) + ' 个智能体人设'}"}, ensure_ascii=False)
    yield f"data: {msg}\n\n"

    # 阶段3：多轮讨论模拟。具体轮次由 round_start 事件逐轮显示。
    phase_msg_discuss = (
        f"Starting multi-round discussion simulation ({rounds} rounds planned)..."
        if is_en else
        f"开始多轮讨论模拟（计划 {rounds} 轮），即将进入第1轮..."
    )
    msg = json.dumps({"type": "phase", "phase": "discussion", "message": phase_msg_discuss}, ensure_ascii=False)
    yield f"data: {msg}\n\n"

    all_speeches = []

    for rnd in range(1, rounds + 1):
        msg = json.dumps({"type": "round_start", "round": rnd}, ensure_ascii=False)
        yield f"data: {msg}\n\n"

        round_speeches = []
        conversation_history = ""
        for prev_speech in all_speeches:
            conversation_history += f"第{prev_speech['round']}轮 - {prev_speech['agent_name']}：{prev_speech['content']}\n\n"

        for agent in agents:
            msg = json.dumps({"type": "agent_speaking", "round": rnd, "agent_id": agent["agent_id"], "agent_name": agent["name"]}, ensure_ascii=False)
            yield f"data: {msg}\n\n"

            speech = await generate_agent_speech(agent, rnd, question, conversation_history, subgraph_context, lang)

            speech_record = {"round": rnd, "agent_id": agent["agent_id"], "agent_name": agent["name"], "content": speech}
            round_speeches.append(speech_record)
            all_speeches.append(speech_record)

            # 推送完整发言内容
            speech_event = {
                "type": "agent_speech",
                "round": rnd,
                "agent_id": agent["agent_id"],
                "agent_name": agent["name"],
                "content": speech,
                "entity_id": agent.get("entity_id", ""),
                "entity_type": agent.get("entity_type", "concept")
            }
            msg = json.dumps(speech_event, ensure_ascii=False)
            yield f"data: {msg}\n\n"

        # 本轮摘要
        msg = json.dumps({"type": "round_summarizing", "round": rnd}, ensure_ascii=False)
        yield f"data: {msg}\n\n"
        round_summary = await generate_round_summary(rnd, question, round_speeches, lang)
        msg = json.dumps({"type": "round_summary", "round": rnd, "summary": round_summary}, ensure_ascii=False)
        yield f"data: {msg}\n\n"

    msg = json.dumps({"type": "phase", "phase": "discussion", "message": f"{rounds} {'rounds of discussion completed' if is_en else '轮讨论模拟完成'}"}, ensure_ascii=False)
    yield f"data: {msg}\n\n"

    # 阶段4：协调者总结
    phase_msg_coord = "Coordinator is summarizing the discussion..." if is_en else "协调者正在总结讨论..."
    msg = json.dumps({"type": "phase", "phase": "coordinator_summary", "message": phase_msg_coord}, ensure_ascii=False)
    yield f"data: {msg}\n\n"

    coordinator_summary = await generate_coordinator_summary(question, all_speeches, rounds, lang)
    msg = json.dumps({"type": "coordinator_summary", "content": coordinator_summary}, ensure_ascii=False)
    yield f"data: {msg}\n\n"

    # 阶段5：最终预测报告
    phase_msg_pred = "Generating final prediction report..." if is_en else "正在生成最终预测报告..."
    msg = json.dumps({"type": "phase", "phase": "prediction", "message": phase_msg_pred}, ensure_ascii=False)
    yield f"data: {msg}\n\n"

    confidence_match = re.search(r'(?:置信度|Confidence)[:：\s]*([0-9.]+)', coordinator_summary)
    confidence = float(confidence_match.group(1)) if confidence_match else 0.75
    if confidence > 1:
        confidence = confidence / 100

    # 尝试中英文两种格式提取预测结论
    conclusion_match = re.search(r'### (?:综合预测|Comprehensive Prediction)(.*?)(?=###|\Z)', coordinator_summary, re.DOTALL)
    conclusion = conclusion_match.group(1).strip() if conclusion_match else coordinator_summary
    conclusion = complete_final_prediction_report(question, conclusion, all_speeches, agents, confidence, lang)

    # 提取关键因素/不确定因素
    key_factors = []
    factors_match = re.search(r'(?:不确定因素|Uncertainties)[^:\n]*[:：\n]\s*(.*?)(?=(?:###|$))', coordinator_summary, re.DOTALL)
    if factors_match:
        factor_text = factors_match.group(1).strip()
        key_factors = [f.strip().lstrip('-* ') for f in factor_text.split('\n') if f.strip() and f.strip().lstrip('-* ')][:6]

    evidence_ids = re.findall(r'\[([a-z_0-9]+)\]', coordinator_summary)
    evidence = list(set(evidence_ids))[:10]

    prediction_data = {
        "type": "prediction",
        "conclusion": conclusion,
        "confidence": confidence,
        "evidence": evidence,
        "key_factors": key_factors,
        "coordinator_summary": coordinator_summary,
        "rounds_count": rounds,
        "speech_count": len(all_speeches),
        "agents": [
            {
                "id": a["agent_id"],
                "name": a["name"],
                "role": a["role"],
                "stance": a.get("stance", "neutral"),
                "expertise": a.get("expertise", ""),
                "entity_type": a.get("entity_type", "")
            }
            for a in agents
        ]
    }
    msg = json.dumps(prediction_data, ensure_ascii=False)
    yield f"data: {msg}\n\n"

    msg = json.dumps({"type": "phase", "phase": "complete", "message": ("Multi-agent simulation completed" if is_en else "多智能体模拟完成")}, ensure_ascii=False)
    yield f"data: {msg}\n\n"


# ============================================================
# API 请求模型
# ============================================================

class ReasonRequest(BaseModel):
    question: str
    mode: str = "reasoning"
    max_nodes: int = 80
    rounds: int = 3
    lang: str = "zh"  # 'zh' or 'en' - 自动检测的问题语言
    llm_config: Optional[Dict[str, Any]] = Field(default=None, alias="model_config")


class SubgraphRequest(BaseModel):
    question: str
    max_nodes: int = 50


class ModelTestRequest(BaseModel):
    provider: str
    api_kind: str = "openai"
    model: str
    base_url: str = ""
    api_key: str


# ============================================================
# FastAPI 应用
# ============================================================

app = FastAPI(title="量子知识图谱推演引擎", version="2.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.api_route("/", methods=["GET", "HEAD"])
async def serve_index():
    """服务主页面"""
    try:
        with open(HTML_FILE_PATH, "r", encoding="utf-8") as f:
            content = f.read()
        return HTMLResponse(content=content)
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail=f"前端文件不存在: {HTML_FILE_PATH}")


def resolve_static_file(static_path: str) -> Tuple[str, str]:
    """把前端静态文件安全地映射到项目目录。"""
    requested = (static_path or "quantum_reasoning.html").lstrip("/")
    if requested.startswith("api/"):
        raise HTTPException(status_code=404, detail="API route not found")

    full_path = os.path.abspath(os.path.join(BASE_DIR, requested))
    if full_path != BASE_DIR and not full_path.startswith(BASE_DIR + os.sep):
        raise HTTPException(status_code=403, detail="Forbidden")
    if not os.path.isfile(full_path):
        raise HTTPException(status_code=404, detail=f"静态文件不存在: {requested}")

    ext = os.path.splitext(full_path)[1].lower()
    return full_path, STATIC_MIME_TYPES.get(ext, "application/octet-stream")


@app.get("/api/stats")
async def get_stats():
    """获取图谱统计信息"""
    if not GRAPH_DATA.get("nodes"):
        raise HTTPException(status_code=503, detail="图谱数据未加载")

    type_counts = Counter(n["type"] for n in GRAPH_DATA["nodes"])
    edge_label_counts = Counter(e.get("label", "(无标签)") for e in EDGES)

    degree_dist = {}
    for nid in NODES:
        deg = len(ADJ[nid])
        bucket = min(deg // 10 * 10, 100)
        key = f"{bucket}-{bucket+9}"
        degree_dist[key] = degree_dist.get(key, 0) + 1

    return {
        "total_nodes": len(GRAPH_DATA["nodes"]),
        "total_edges": len(EDGES),
        "entity_types": {TYPE_LABELS.get(k, k): v for k, v in type_counts.most_common()},
        "top_relations": dict(edge_label_counts.most_common(20)),
        "degree_distribution": degree_dist,
        "graph_metadata": GRAPH_DATA.get("metadata", {})
    }


@app.get("/api/models/providers")
async def get_model_providers():
    """返回前端可选的大模型厂商预设。"""
    return {"providers": MODEL_PROVIDER_PRESETS}


@app.post("/api/models/test")
async def test_model(req: ModelTestRequest):
    """测试用户模型配置，测试通过后前端才允许用于推演。"""
    try:
        raw_config = req.model_dump() if hasattr(req, "model_dump") else req.dict()
        config = require_model_config(raw_config)
        test_messages = [
            {"role": "system", "content": "You are a connection test assistant. Reply only with OK."},
            {"role": "user", "content": "Reply OK if this model connection works."}
        ]
        token = ACTIVE_MODEL_CONFIG.set(config)
        started = time.time()
        try:
            result = await call_deepseek(test_messages, temperature=0, max_tokens=128)
        finally:
            ACTIVE_MODEL_CONFIG.reset(token)

        if not result or not result.strip():
            raise HTTPException(status_code=502, detail="模型没有返回有效内容")
        return {
            "ok": True,
            "provider": config["provider"],
            "provider_label": config["provider_label"],
            "model": config["model"],
            "base_url": config["base_url"],
            "latency_ms": int((time.time() - started) * 1000),
            "sample": result.strip()[:120]
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"模型测试调用异常: {type(e).__name__}: {str(e)}")


@app.post("/api/subgraph")
async def get_relevant_subgraph(req: SubgraphRequest):
    """获取与问题相关的子图"""
    if not GRAPH_DATA.get("nodes"):
        raise HTTPException(status_code=503, detail="图谱数据未加载")

    seeds = find_seed_nodes(req.question, top_k=15)
    if not seeds:
        seeds = sorted(GRAPH_DATA["nodes"], key=lambda n: n.get("properties", {}).get("importance", 0), reverse=True)[:5]

    seed_ids = [n["id"] for n in seeds]
    subgraph = get_subgraph(seed_ids, max_nodes=req.max_nodes)

    seed_id_set = set(seed_ids)
    for n in subgraph["nodes"]:
        n["is_seed"] = n["id"] in seed_id_set

    return {
        "seed_nodes": [{"id": n["id"], "label": n["label"], "type": n["type"]} for n in seeds],
        "subgraph": subgraph,
        "stats": {
            "seed_count": len(seeds),
            "subgraph_nodes": len(subgraph["nodes"]),
            "subgraph_edges": len(subgraph["edges"])
        }
    }


@app.post("/api/reason")
async def reason(req: ReasonRequest):
    """推理推演（SSE流式输出）"""
    if not GRAPH_DATA.get("nodes"):
        raise HTTPException(status_code=503, detail="图谱数据未加载")

    return StreamingResponse(
        stream_with_model_config(reasoning_mode(req.question, max_nodes=req.max_nodes, lang=req.lang), req.llm_config, "reason"),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"}
    )


@app.post("/api/simulate")
async def simulate(req: ReasonRequest):
    """多智能体模拟（SSE流式输出）"""
    if not GRAPH_DATA.get("nodes"):
        raise HTTPException(status_code=503, detail="图谱数据未加载")

    return StreamingResponse(
        stream_with_model_config(simulation_mode(req.question, rounds=req.rounds, max_nodes=req.max_nodes, lang=req.lang), req.llm_config, "simulate"),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"}
    )


@app.get("/api/graph/full")
async def get_full_graph(limit_nodes: int = 500, limit_edges: int = 2000):
    """获取完整图谱数据"""
    nodes = GRAPH_DATA.get("nodes", [])[:limit_nodes]
    edges = GRAPH_DATA.get("edges", [])[:limit_edges]
    return {
        "nodes": nodes,
        "edges": edges,
        "truncated": len(GRAPH_DATA.get("nodes", [])) > limit_nodes or len(GRAPH_DATA.get("edges", [])) > limit_edges,
        "total_nodes": len(GRAPH_DATA.get("nodes", [])),
        "total_edges": len(GRAPH_DATA.get("edges", []))
    }


@app.get("/api/graph/neighborhood/{node_id}")
async def get_node_neighborhood_api(node_id: str, max_hops: int = 2):
    """获取某个节点的邻域子图"""
    if node_id not in NODES:
        raise HTTPException(status_code=404, detail=f"节点不存在: {node_id}")
    subgraph = get_node_neighborhood(node_id, max_hops=max_hops)
    return subgraph


@app.get("/api/health")
async def health_check():
    """健康检查"""
    return {
        "status": "ok",
        "base_dir": BASE_DIR,
        "graph_loaded": len(GRAPH_DATA.get("nodes", [])) > 0,
        "node_count": len(GRAPH_DATA.get("nodes", [])),
        "edge_count": len(GRAPH_DATA.get("edges", [])),
        "default_model_ready": bool(DEEPSEEK_API_KEY)
    }


# ============================================================
# IMA 知识库引用
# ============================================================

# IMA知识库映射
IMA_KNOWLEDGE_BASES = {
    "7412047272751870": {"name": "量子计算学习材料", "count": 96},
    "7425064727240963": {"name": "量子产业研究报告", "count": 106},
    "7411640966326739": {"name": "金贻荣老师的直播课件", "count": 79},
    "7408677623443647": {"name": "量子计算", "count": 44},
    "7441506466037417": {"name": "李博士的项目", "count": 111},
}

class IMAReferenceRequest(BaseModel):
    question: str
    subgraph_nodes: List[str] = []


@app.post("/api/ima/references")
async def get_ima_references(req: IMAReferenceRequest):
    """获取IMA知识库相关引用"""
    references = []

    # 从子图节点提取source_knowledge_bases
    kb_hits = defaultdict(list)
    for node_id in req.subgraph_nodes:
        if node_id in NODES:
            node = NODES[node_id]
            source_kbs = node.get("properties", {}).get("source_knowledge_bases", [])
            for kb_id in source_kbs:
                if kb_id in IMA_KNOWLEDGE_BASES:
                    kb_hits[kb_id].append(node["label"])

    # 基于问题关键词匹配知识库
    q_lower = req.question.lower()
    for kb_id, kb_info in IMA_KNOWLEDGE_BASES.items():
        score = 0
        # 关键词匹配
        for term in ["量子", "计算", "超导", "离子", "纠错", "算法", "比特", "产业", "芯片", "公司"]:
            if term in q_lower and term in kb_info["name"].lower():
                score += 2
        if kb_id in kb_hits:
            score += len(kb_hits[kb_id]) * 3
        if score > 0:
            references.append({
                "kb_id": kb_id,
                "kb_name": kb_info["name"],
                "title": f"《{kb_info['name']}》知识库",
                "doc_count": kb_info["count"],
                "related_nodes": kb_hits.get(kb_id, [])[:5],
                "content": f"该知识库包含{kb_info['count']}条量子计算相关文档，涵盖{', '.join(kb_hits.get(kb_id, ['综合知识'])[:3])}等领域。" if score > 3 else None
            })

    # 用DeepSeek搜索获取更具体的引用
    try:
        search_prompt = f"""基于问题「{req.question}」，从以下5个量子计算知识库中列出最相关的文档标题和简要说明：
1. 量子计算学习材料 (96条) - 包含量子计算基础教程、算法教程等
2. 量子产业研究报告 (106条) - 包含量子产业分析、市场报告、公司分析等
3. 金贻荣老师的直播课件 (79条) - 包含量子计算教学课件、实验指导等
4. 量子计算 (44条) - 包含量子计算核心概念、原理介绍等
5. 李博士的项目 (111条) - 包含量子研究项目报告、实验数据等

请以JSON数组格式返回，每项包含: kb_name, title, relevance(1-5), brief(50字以内简介)。最多返回8条。只返回JSON数组。"""

        response = await call_deepseek(
            [{"role": "user", "content": search_prompt}],
            temperature=0.3, max_tokens=1024
        )
        json_match = re.search(r'\[.*\]', response, re.DOTALL)
        if json_match:
            llm_refs = json.loads(json_match.group())
            for ref in llm_refs[:8]:
                kb_name = ref.get("kb_name", "")
                # 匹配到已有的知识库
                matched = False
                for existing in references:
                    if existing["kb_name"] == kb_name:
                        matched = True
                        break
                if not matched:
                    # 查找知识库ID
                    kb_id = ""
                    for kid, kinfo in IMA_KNOWLEDGE_BASES.items():
                        if kinfo["name"] == kb_name:
                            kb_id = kid
                            break
                    references.append({
                        "kb_id": kb_id,
                        "kb_name": kb_name,
                        "title": ref.get("title", ""),
                        "doc_count": IMA_KNOWLEDGE_BASES.get(kb_id, {}).get("count", 0),
                        "related_nodes": [],
                        "content": ref.get("brief", f"来自「{kb_name}」知识库的相关文档")
                    })
    except Exception as e:
        print(f"[IMA引用] LLM搜索失败: {e}")

    # 按相关性排序
    references.sort(key=lambda x: len(x.get("related_nodes", [])), reverse=True)

    return {"references": references[:12]}


# ============================================================
# 报告导出
# ============================================================

class ExportRequest(BaseModel):
    question: str
    mode: str = "reasoning"
    steps: List[Dict[str, Any]] = []
    prediction: Dict[str, Any] = {}
    ima_references: List[Dict[str, Any]] = []
    lang: str = "zh"
    format: str = "html"
    export_dir: str = ""
    simulation_data: Optional[Dict[str, Any]] = None  # 模拟模式中间过程数据


def clean_report_text(value: Any) -> str:
    """把模型 Markdown 清洗成正式报告里的纯文本。"""
    text = str(value or "").replace("\r\n", "\n").replace("\r", "\n")
    lines = []
    in_code = False
    for raw_line in text.split("\n"):
        line = raw_line.strip()
        if line.startswith("```"):
            in_code = not in_code
            continue
        line = re.sub(r"^\s*>+\s*", "", line)
        line = re.sub(r"^\s*#{1,6}\s*", "", line)
        line = re.sub(r"^\s*[-*+]\s+", "", line)
        line = re.sub(r"^\s*\d+[.)]\s+", "", line)
        line = line.replace("**", "").replace("__", "")
        line = re.sub(r"`([^`]*)`", r"\1", line)
        line = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", line)
        lines.append(line if in_code else line.strip())
    cleaned = "\n".join(lines)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()
    return cleaned


def normalize_doc_count(value: Any) -> int:
    try:
        count = int(float(value or 0))
    except (TypeError, ValueError):
        count = 0
    return max(1, count)


def normalize_ima_references(refs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """规范导出用引用，确保文档数至少为1，并保留章节/内容片段。"""
    normalized = []
    seen = set()
    for ref in refs or []:
        title = clean_report_text(ref.get("title", ""))
        if not title:
            continue
        kb_name = clean_report_text(ref.get("kb_name", "")) or "IMA知识库"
        key = (kb_name, title)
        if key in seen:
            continue
        seen.add(key)
        related_nodes = ref.get("related_nodes") or ref.get("nodes") or []
        if isinstance(related_nodes, str):
            related_nodes = [related_nodes]
        sections = ref.get("sections") or []
        if isinstance(sections, str):
            sections = [sections]
        content = ref.get("content") or "\n\n".join(sections)
        normalized.append({
            "title": title,
            "kb_name": kb_name,
            "doc_count": normalize_doc_count(ref.get("doc_count", 1)),
            "related_nodes": [clean_report_text(n) for n in related_nodes if clean_report_text(n)][:12],
            "content": clean_report_text(content),
        })
    return normalized[:30]


def build_report_html(req: ExportRequest) -> str:
    """构建HTML报告"""
    import html as html_module

    def escape_html(s: str) -> str:
        return html_module.escape(s or "")

    def report_html_text(value: Any) -> str:
        return escape_html(clean_report_text(value)).replace("\n", "<br>")

    is_zh = req.lang == "zh"
    is_en = not is_zh
    sim_data = req.simulation_data

    # 标题
    title = f"量子计算知识图谱推演报告" if is_zh else "Quantum Computing Knowledge Graph Reasoning Report"
    subtitle = clean_report_text(req.question)
    date_str = time.strftime("%Y-%m-%d %H:%M")

    # === 推理模式步骤 ===
    steps_html = ""
    for step in req.steps:
        step_title = clean_report_text(step.get("title", ""))
        step_content = report_html_text(step.get("content", ""))
        steps_html += f"""
        <div style="margin-bottom:20px;page-break-inside:avoid">
          <h3 style="color:#3b82f6;border-bottom:2px solid #3b82f6;padding-bottom:6px">{escape_html(step_title)}</h3>
          <div style="padding:12px;background:#f8fafc;border-radius:6px;line-height:1.8;font-size:14px">{step_content}</div>
        </div>"""

    # === 模拟模式中间过程 ===
    simulation_html = ""
    if sim_data and req.mode == "simulation":
        # 阶段记录
        if sim_data.get("phaseMessages"):
            phase_html_parts = []
            for pm in sim_data["phaseMessages"]:
                phase_html_parts.append(f'<li style="padding:4px 0;color:#60A5FA">{report_html_text(pm.get("message", ""))}</li>')
            simulation_html += f"""
        <div style="margin-bottom:20px;page-break-inside:avoid">
          <h3 style="color:#8b5cf6;border-bottom:2px solid #8b5cf6;padding-bottom:6px">{'📋 执行阶段' if is_zh else '📋 Execution Phases'}</h3>
          <ul style="padding-left:20px;line-height:1.8">{"".join(phase_html_parts)}</ul>
        </div>"""

        # 智能体人设列表
        agents_list = sim_data.get("agents", [])
        if agents_list:
            agent_cards = ""
            for a in agents_list:
                stance_color = {"optimistic": "#34D399", "pessimistic": "#F87171", "neutral": "#60A5FA"}.get(a.get("stance", "neutral"), "#60A5FA")
                stance_label = {"optimistic": ("乐观" if is_zh else "Optimistic"), "pessimistic": ("悲观" if is_zh else "Pessimistic"), "neutral": ("中立" if is_zh else "Neutral")}.get(a.get("stance", "neutral"), "中立" if is_zh else "Neutral")
                persona_label = "人设" if is_zh else "Persona"
                persona_html = (
                    f'<div style="font-size:12px;color:#64748b;margin-top:4px"><b>{persona_label}:</b> {report_html_text(a.get("persona", ""))[:300]}</div>'
                    if a.get("persona") else ""
                )
                agent_cards += f"""
            <div style="margin:10px 0;padding:14px;background:#f8fafc;border-radius:8px;border-left:4px solid {stance_color}">
              <div style="font-weight:700;font-size:15px;color:#1e293b">{report_html_text(a.get('name', ''))}</div>
              <div style="display:flex;gap:8px;margin:4px 0;flex-wrap:wrap">
                <span style="font-size:11px;background:#e0e7ff;color:#4338ca;padding:2px 8px;border-radius:4px">{report_html_text(a.get('role', ''))}</span>
                <span style="font-size:11px;background:{stance_color}22;color:{stance_color};padding:2px 8px;border-radius:4px">{stance_label}</span>
              </div>
              <div style="font-size:12px;color:#64748b;margin-top:4px"><b>{'专业领域' if is_zh else 'Expertise'}:</b> {report_html_text(a.get('expertise', ''))}</div>
              {persona_html}
            </div>"""
            simulation_html += f"""
        <div style="margin-bottom:20px;page-break-inside:avoid">
          <h3 style="color:#8b5cf6;border-bottom:2px solid #8b5cf6;padding-bottom:6px">{'👥 参与智能体' if is_zh else '👥 Participating Agents'} ({len(agents_list)})</h3>
          {agent_cards}
        </div>"""

        # 多轮讨论详情
        rounds_list = sim_data.get("rounds", [])
        if rounds_list:
            for rnd in rounds_list:
                rnd_num = rnd.get("round", "?")
                speeches_html = ""
                for sp in rnd.get("speeches", []):
                    agent_name = sp.get("agent_name", "")
                    content = report_html_text(sp.get("content", ""))
                    speeches_html += f"""
              <div style="margin:10px 0;padding:12px;background:#fafafa;border-radius:6px">
                <div style="font-weight:700;font-size:13px;color:#374151;margin-bottom:4px">{report_html_text(agent_name)}</div>
                <div style="font-size:13px;color:#4b5563;line-height:1.8">{content}</div>
              </div>"""

                round_summary = clean_report_text(rnd.get("summary", ""))
                summary_html = ""
                if round_summary:
                    summary_html = f"""
              <div style="margin:8px 0;padding:10px;background:#f0f9ff;border-radius:6px;border:1px solid #bae6fd">
                <div style="font-weight:700;font-size:12px;color:#0369a1">{'📋 轮次摘要' if is_zh else '📋 Round Summary'}</div>
                <div style="font-size:12px;color:#0c4a6e;line-height:1.7;margin-top:4px">{report_html_text(round_summary)}</div>
              </div>"""

                simulation_html += f"""
        <div style="margin-bottom:24px;page-break-inside:avoid">
          <h3 style="color:#06b6d4;border-bottom:2px solid #06b6d4;padding-bottom:6px">{'💬 第' + str(rnd_num) + '轮讨论' if is_zh else '💬 Round ' + str(rnd_num) + ' Discussion'}</h3>
          {speeches_html}{summary_html}
        </div>"""

        # 协调者总结
        coord_sum = clean_report_text(sim_data.get("coordinatorSummary", ""))
        if coord_sum:
            simulation_html += f"""
        <div style="margin-bottom:20px;page-break-inside:avoid">
          <h3 style="color:#c084fc;border-bottom:2px solid #c084fc;padding-bottom:6px">{'🎯 协调者总结' if is_zh else '🎯 Coordinator Summary'}</h3>
          <div style="padding:16px;background:#faf5ff;border-radius:8px;line-height:1.9;font-size:14px">{report_html_text(coord_sum)}</div>
        </div>"""

    # 预测结论
    pred = req.prediction
    confidence = pred.get("confidence", 0.7)
    conf_pct = int(confidence * 100)
    conclusion = report_html_text(pred.get("conclusion", ""))
    evidence = pred.get("evidence", [])
    key_factors = pred.get("key_factors", [])

    evidence_rows = ""
    for e in evidence:
        evidence_rows += f"<tr><td style='padding:6px 10px;border:1px solid #e2e8f0'>{report_html_text(e)}</td></tr>"

    factors_rows = ""
    for f in key_factors:
        factors_rows += f"<tr><td style='padding:6px 10px;border:1px solid #e2e8f0'>{report_html_text(f)}</td></tr>"

    # IMA引用
    ima_html = ""
    for i, ref in enumerate(normalize_ima_references(req.ima_references), start=1):
        related_nodes = "、".join(ref.get("related_nodes", []))
        related_html = f"<div style='font-size:12px;color:#059669;margin-top:6px'>{'关联节点' if is_zh else 'Related nodes'}: {escape_html(related_nodes)}</div>" if related_nodes else ""
        content_html = f"<div style='margin-top:8px;padding:10px;background:#f8fafc;border-radius:6px;line-height:1.7;font-size:12px;color:#475569'>{report_html_text(ref.get('content', ''))}</div>" if ref.get("content") else ""
        ima_html += f"""
        <div style="margin:10px 0;padding:12px;border:1px solid #e2e8f0;border-radius:8px;page-break-inside:avoid">
          <div style="font-weight:700;color:#1e293b">{i}. {escape_html(ref.get("title", ""))}</div>
          <div style="font-size:12px;color:#64748b;margin-top:4px">{'知识库' if is_zh else 'Knowledge Base'}: {escape_html(ref.get("kb_name", ""))} ｜ {'文档数' if is_zh else 'Docs'}: {normalize_doc_count(ref.get("doc_count", 1))}</div>
          {related_html}
          {content_html}
        </div>"""

    mode_label = "逐步推理" if req.mode == "reasoning" else "多智能体模拟"
    if not is_zh:
        mode_label = "Step-by-step Reasoning" if req.mode == "reasoning" else "Multi-Agent Simulation"

    simulation_section_html = ""
    if simulation_html:
        simulation_title = "👥 多智能体推演过程" if is_zh else "👥 Multi-Agent Simulation Process"
        simulation_section_html = f'<div class="section"><h2>{simulation_title}</h2>{simulation_html}</div>'

    evidence_table_html = ""
    if evidence_rows:
        evidence_title = "支撑证据" if is_zh else "Supporting Evidence"
        evidence_table_html = f"<table><tr><th>{evidence_title}</th></tr>{evidence_rows}</table>"

    factors_table_html = ""
    if factors_rows:
        factors_title = "关键因素" if is_zh else "Key Factors"
        factors_table_html = f"<table><tr><th>{factors_title}</th></tr>{factors_rows}</table>"

    ima_section_html = ""
    if ima_html:
        ima_title = "相关文献列表" if is_zh else "Related Literature"
        ima_section_html = f"<div class='section'><h2>📚 {ima_title}</h2>{ima_html}</div>"

    conf_color = "#10b981" if conf_pct > 70 else "#f59e0b" if conf_pct > 40 else "#ef4444"

    html = f"""<!DOCTYPE html>
<html lang="{'zh-CN' if is_zh else 'en'}">
<head>
<meta charset="UTF-8">
<title>{title}</title>
<style>
  @page {{ margin: 2cm; }}
  body {{ font-family: -apple-system, "PingFang SC", "Microsoft YaHei", "Helvetica Neue", Arial, sans-serif; color: #1a1a2e; max-width: 900px; margin: 0 auto; padding: 40px; line-height: 1.8; }}
  h1 {{ color: #1e40af; font-size: 28px; border-bottom: 3px solid #3b82f6; padding-bottom: 10px; }}
  h2 {{ color: #374151; font-size: 20px; margin-top: 30px; }}
  table {{ width: 100%; border-collapse: collapse; margin: 12px 0; }}
  th {{ background: #3b82f6; color: white; padding: 8px 12px; text-align: left; font-size: 13px; }}
  td {{ padding: 6px 12px; border: 1px solid #e2e8f0; font-size: 13px; }}
  .cover {{ text-align: center; padding: 60px 0 40px; }}
  .cover h1 {{ border: none; font-size: 36px; }}
  .meta {{ color: #6b7280; font-size: 14px; margin: 6px 0; }}
  .conf-bar {{ height: 12px; background: #e5e7eb; border-radius: 6px; overflow: hidden; }}
  .conf-fill {{ height: 100%; border-radius: 6px; }}
  .section {{ margin-top: 30px; page-break-inside: avoid; }}
</style>
</head>
<body>
<div class="cover">
  <h1>{title}</h1>
  <div class="meta">{'问题' if is_zh else 'Question'}: {escape_html(subtitle)}</div>
  <div class="meta">{'模式' if is_zh else 'Mode'}: {mode_label}</div>
  <div class="meta">{'日期' if is_zh else 'Date'}: {date_str}</div>
  <div class="meta">{'图谱规模' if is_zh else 'Graph Scale'}: 308 {'节点' if is_zh else 'nodes'}, 1014 {'关系' if is_zh else 'edges'}</div>
</div>

<div class="section">
  <h2>{'📊 推理过程' if is_zh else '📊 Reasoning Process'}</h2>
  {steps_html}
</div>

	{simulation_section_html}

<div class="section">
  <h2>{'🔮 预测结论' if is_zh else '🔮 Prediction Conclusion'}</h2>
  <div style="padding:16px;background:#f0f9ff;border:1px solid #93c5fd;border-radius:8px;margin-bottom:16px">
    {conclusion}
  </div>
  <table>
    <tr><th colspan="2">{'置信度评估' if is_zh else 'Confidence Assessment'}</th></tr>
    <tr>
      <td style="width:120px;font-weight:600">{'置信度' if is_zh else 'Confidence'}</td>
      <td>
        <div style="display:flex;align-items:center;gap:10px">
	          <div class="conf-bar" style="flex:1"><div class="conf-fill" style="width:{conf_pct}%;background:{conf_color}"></div></div>
          <span style="font-weight:700">{conf_pct}%</span>
        </div>
      </td>
    </tr>
  </table>
	  {evidence_table_html}
	  {factors_table_html}
	</div>
	
	{ima_section_html}

<div style="margin-top:40px;padding-top:20px;border-top:1px solid #e2e8f0;color:#9ca3af;font-size:11px;text-align:center">
  {'本报告由量子知识图谱推演引擎自动生成' if is_zh else 'This report was auto-generated by Quantum Knowledge Graph Reasoning Engine'} | {date_str}
</div>
</body>
</html>"""
    return html


def build_report_docx(req: ExportRequest) -> bytes:
    """构建Word报告"""
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    is_zh = req.lang == "zh"
    is_en = not is_zh
    sim_data = req.simulation_data
    doc = Document()

    # 样式设置
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei' if is_zh else 'Arial'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    # 标题
    title = doc.add_heading('量子计算知识图谱推演报告' if is_zh else 'Quantum Computing Knowledge Graph Reasoning Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 元信息
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_para.add_run(f"{'问题' if is_zh else 'Question'}: {clean_report_text(req.question)}\n").bold = True
    meta_para.add_run(f"{'模式' if is_zh else 'Mode'}: {'逐步推理' if is_zh and req.mode == 'reasoning' else '多智能体模拟' if is_zh else req.mode}\n")
    meta_para.add_run(f"{'日期' if is_zh else 'Date'}: {time.strftime('%Y-%m-%d %H:%M')}\n")
    meta_para.add_run(f"{'图谱规模' if is_zh else 'Graph Scale'}: 308 {'节点' if is_zh else 'nodes'}, 1014 {'关系' if is_zh else 'edges'}")

    doc.add_page_break()

    # 推理过程
    doc.add_heading('📊 ' + ('推理过程' if is_zh else 'Reasoning Process'), 1)
    for step in req.steps:
        doc.add_heading(clean_report_text(step.get("title", "")), level=2)
        content = step.get("content", "")
        for line in content.split("\n"):
            if line.strip():
                p = doc.add_paragraph(clean_report_text(line))
                p.paragraph_format.space_after = Pt(4)

    # 模拟模式中间过程
    if sim_data and req.mode == "simulation":
        # 执行阶段
        phase_msgs = sim_data.get("phaseMessages", [])
        if phase_msgs:
            doc.add_heading('📋 ' + ('执行阶段' if is_zh else 'Execution Phases'), 1)
            for pm in phase_msgs:
                doc.add_paragraph(clean_report_text(pm.get("message", "")))

        # 参与智能体
        agents_list = sim_data.get("agents", [])
        if agents_list:
            doc.add_heading('👥 ' + ('参与智能体' if is_zh else 'Participating Agents') + f" ({len(agents_list)})", 1)
            for a in agents_list:
                p = doc.add_paragraph()
                run_name = p.add_run(clean_report_text(a.get("name", "")))
                run_name.bold = True
                run_name.font.size = Pt(13)
                p.add_run(f" {clean_report_text(a.get('role', ''))} | {clean_report_text(a.get('stance', 'neutral'))}")
                if a.get("expertise"):
                    exp_p = doc.add_paragraph()
                    exp_p.add_run(('专业领域: ' if is_zh else 'Expertise: ')).italic = True
                    exp_p.add_run(clean_report_text(a["expertise"]))
                if a.get("persona"):
                    pera_p = doc.add_paragraph()
                    pera_p.add_run(('人设: ' if is_zh else 'Persona: ')).italic = True
                    pera_p.add_run(clean_report_text(a["persona"])[:300])

        # 多轮讨论
        rounds_list = sim_data.get("rounds", [])
        if rounds_list:
            for rnd in rounds_list:
                rnd_num = rnd.get("round", "?")
                doc.add_heading('💬 ' + (('第' + str(rnd_num) + '轮讨论') if is_zh else ('Round ' + str(rnd_num) + ' Discussion')), level=2)
                for sp in rnd.get("speeches", []):
                    sp_p = doc.add_paragraph()
                    sp_run = sp_p.add_run(clean_report_text(sp.get('agent_name', '')))
                    sp_run.bold = True
                    sp_content = sp.get("content", "")
                    for line in sp_content.split("\n"):
                        if line.strip():
                            lc_p = doc.add_paragraph(clean_report_text(line))
                            lc_p.paragraph_format.left_indent = Inches(0.3)

                rnd_summary = rnd.get("summary", "")
                if rnd_summary:
                    sum_p = doc.add_paragraph()
                    sum_run = sum_p.add_run(('📋 ' + ('轮次摘要' if is_zh else 'Round Summary')) + ":")
                    sum_run.bold = True
                    sum_run.font.color.rgb = RGBColor(3, 105, 161)
                    doc.add_paragraph(clean_report_text(rnd_summary))

        # 协调者总结
        coord_sum = sim_data.get("coordinatorSummary", "")
        if coord_sum:
            doc.add_heading('🎯 ' + ('协调者总结' if is_zh else 'Coordinator Summary'), 1)
            for line in coord_sum.split("\n"):
                if line.strip():
                    p = doc.add_paragraph(clean_report_text(line))

    # 预测结论
    doc.add_heading('🔮 ' + ('预测结论' if is_zh else 'Prediction Conclusion'), 1)
    pred = req.prediction
    conclusion = clean_report_text(pred.get("conclusion", ""))
    for line in conclusion.split("\n"):
        if line.strip():
            doc.add_paragraph(clean_report_text(line))

    # 置信度表
    confidence = pred.get("confidence", 0.7)
    conf_pct = int(confidence * 100)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = '置信度' if is_zh else 'Confidence'
    hdr[1].text = f"{conf_pct}%"
    doc.add_paragraph()

    # 证据表
    evidence = pred.get("evidence", [])
    if evidence:
        doc.add_heading('支撑证据' if is_zh else 'Supporting Evidence', level=2)
        ev_table = doc.add_table(rows=1, cols=1)
        ev_table.style = 'Light Grid Accent 1'
        ev_table.rows[0].cells[0].text = '证据' if is_zh else 'Evidence'
        for e in evidence:
            row = ev_table.add_row().cells
            row[0].text = clean_report_text(e)

    # IMA引用
    ima_refs = normalize_ima_references(req.ima_references)
    if ima_refs:
        doc.add_heading('📚 ' + ('相关文献列表' if is_zh else 'Related Literature'), 1)
        ima_table = doc.add_table(rows=1, cols=3)
        ima_table.style = 'Light Grid Accent 1'
        hdr = ima_table.rows[0].cells
        hdr[0].text = '文档标题' if is_zh else 'Title'
        hdr[1].text = '知识库' if is_zh else 'Knowledge Base'
        hdr[2].text = '文档数' if is_zh else 'Docs'
        for ref in ima_refs:
            row = ima_table.add_row().cells
            row[0].text = ref.get("title", "")
            row[1].text = ref.get("kb_name", "")
            row[2].text = str(normalize_doc_count(ref.get("doc_count", 1)))
            if ref.get("related_nodes"):
                doc.add_paragraph(('关联节点: ' if is_zh else 'Related nodes: ') + '、'.join(ref["related_nodes"]))
            if ref.get("content"):
                doc.add_paragraph(clean_report_text(ref["content"]))

    # 页脚
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('— ' + ('本报告由量子知识图谱推演引擎自动生成' if is_zh else 'Auto-generated by Quantum KG Reasoning Engine') + ' —').font.color.rgb = RGBColor(150, 150, 150)

    # 保存到字节
    import io
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_report_pdf(req: ExportRequest) -> bytes:
    """使用 ReportLab 构建不依赖系统图形库的 PDF 报告。"""
    import io
    import html as html_module
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfbase import pdfmetrics

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    font_name = "STSong-Light"
    is_zh = req.lang == "zh"

    def esc(value: Any) -> str:
        return html_module.escape(clean_report_text(value)).replace("\n", "<br/>")

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=1.6 * cm,
        leftMargin=1.6 * cm,
        topMargin=1.6 * cm,
        bottomMargin=1.6 * cm,
        title="量子计算知识图谱推演报告" if is_zh else "Quantum Computing Knowledge Graph Reasoning Report",
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="QTitle", parent=styles["Title"], fontName=font_name, fontSize=22, leading=28, alignment=TA_CENTER, textColor=colors.HexColor("#1e40af")))
    styles.add(ParagraphStyle(name="QHeading", parent=styles["Heading2"], fontName=font_name, fontSize=15, leading=20, textColor=colors.HexColor("#1e40af"), spaceBefore=10, spaceAfter=8))
    styles.add(ParagraphStyle(name="QBody", parent=styles["BodyText"], fontName=font_name, fontSize=10.5, leading=16, spaceAfter=6))
    styles.add(ParagraphStyle(name="QMeta", parent=styles["BodyText"], fontName=font_name, fontSize=10, leading=15, alignment=TA_CENTER, textColor=colors.HexColor("#64748b")))

    story = []
    title = "量子计算知识图谱推演报告" if is_zh else "Quantum Computing Knowledge Graph Reasoning Report"
    mode_label = "逐步推理" if req.mode == "reasoning" and is_zh else "多智能体模拟" if is_zh else req.mode

    story.append(Paragraph(title, styles["QTitle"]))
    story.append(Spacer(1, 0.35 * cm))
    story.append(Paragraph(f"{'问题' if is_zh else 'Question'}: {esc(req.question)}", styles["QMeta"]))
    story.append(Paragraph(f"{'模式' if is_zh else 'Mode'}: {esc(mode_label)}", styles["QMeta"]))
    story.append(Paragraph(f"{'日期' if is_zh else 'Date'}: {time.strftime('%Y-%m-%d %H:%M')}", styles["QMeta"]))
    story.append(Paragraph(f"{'图谱规模' if is_zh else 'Graph Scale'}: 308 {'节点' if is_zh else 'nodes'}, 1014 {'关系' if is_zh else 'edges'}", styles["QMeta"]))
    story.append(PageBreak())

    if req.steps:
        story.append(Paragraph("推理过程" if is_zh else "Reasoning Process", styles["QHeading"]))
        for step in req.steps:
            story.append(Paragraph(esc(step.get("title", "")), styles["QHeading"]))
            story.append(Paragraph(esc(step.get("content", "")), styles["QBody"]))

    sim_data = req.simulation_data or {}
    if sim_data and req.mode == "simulation":
        story.append(Paragraph("多智能体推演过程" if is_zh else "Multi-Agent Simulation Process", styles["QHeading"]))
        for pm in sim_data.get("phaseMessages", []):
            story.append(Paragraph(esc(pm.get("message", "")), styles["QBody"]))
        for agent in sim_data.get("agents", []):
            text = f"<b>{esc(agent.get('name', ''))}</b> - {esc(agent.get('role', ''))} / {esc(agent.get('stance', 'neutral'))}<br/>{esc(agent.get('expertise', ''))}"
            story.append(Paragraph(text, styles["QBody"]))
        for rnd in sim_data.get("rounds", []):
            story.append(Paragraph((f"第{rnd.get('round', '?')}轮讨论" if is_zh else f"Round {rnd.get('round', '?')} Discussion"), styles["QHeading"]))
            for speech in rnd.get("speeches", []):
                story.append(Paragraph(f"<b>{esc(speech.get('agent_name', ''))}</b>: {esc(speech.get('content', ''))}", styles["QBody"]))
            if rnd.get("summary"):
                story.append(Paragraph(f"<b>{'轮次摘要' if is_zh else 'Round Summary'}:</b> {esc(rnd.get('summary', ''))}", styles["QBody"]))
        if sim_data.get("coordinatorSummary"):
            story.append(Paragraph("协调者总结" if is_zh else "Coordinator Summary", styles["QHeading"]))
            story.append(Paragraph(esc(sim_data.get("coordinatorSummary", "")), styles["QBody"]))

    pred = req.prediction or {}
    story.append(Paragraph("预测结论" if is_zh else "Prediction Conclusion", styles["QHeading"]))
    story.append(Paragraph(esc(pred.get("conclusion", "")), styles["QBody"]))
    confidence = int(float(pred.get("confidence", 0.7) or 0.7) * 100)
    story.append(Paragraph(f"{'置信度' if is_zh else 'Confidence'}: {confidence}%", styles["QBody"]))

    evidence = pred.get("evidence", []) or []
    if evidence:
        data = [[Paragraph("支撑证据" if is_zh else "Supporting Evidence", styles["QBody"])]]
        data += [[Paragraph(esc(item), styles["QBody"])] for item in evidence]
        table = Table(data, colWidths=[17 * cm])
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dbeafe")),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#cbd5e1")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(table)
        story.append(Spacer(1, 0.25 * cm))

    ima_refs = normalize_ima_references(req.ima_references)
    if ima_refs:
        story.append(Paragraph("相关文献列表" if is_zh else "Related Literature", styles["QHeading"]))
        data = [[
            Paragraph("标题" if is_zh else "Title", styles["QBody"]),
            Paragraph("知识库" if is_zh else "Knowledge Base", styles["QBody"]),
            Paragraph("文档数" if is_zh else "Docs", styles["QBody"])
        ]]
        for ref in ima_refs:
            data.append([
                Paragraph(esc(ref.get("title", "")), styles["QBody"]),
                Paragraph(esc(ref.get("kb_name", "")), styles["QBody"]),
                Paragraph(str(normalize_doc_count(ref.get("doc_count", 1))), styles["QBody"])
            ])
        table = Table(data, colWidths=[9 * cm, 6 * cm, 2 * cm])
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dbeafe")),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#cbd5e1")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(table)
        story.append(Spacer(1, 0.2 * cm))
        for ref in ima_refs:
            if ref.get("related_nodes"):
                story.append(Paragraph(f"{'关联节点' if is_zh else 'Related nodes'}: {esc('、'.join(ref['related_nodes']))}", styles["QBody"]))
            if ref.get("content"):
                story.append(Paragraph(esc(ref.get("content", "")), styles["QBody"]))
                story.append(Spacer(1, 0.15 * cm))

    story.append(Spacer(1, 0.6 * cm))
    story.append(Paragraph("本报告由量子知识图谱推演引擎自动生成" if is_zh else "Auto-generated by Quantum Knowledge Graph Reasoning Engine", styles["QMeta"]))
    doc.build(story)
    return buffer.getvalue()


@app.post("/api/export/report")
async def export_report(req: ExportRequest):
    """导出推演报告"""
    try:
        export_dir = os.path.abspath(os.path.expanduser(req.export_dir or DEFAULT_EXPORT_DIR))
        os.makedirs(export_dir, exist_ok=True)

        timestamp = time.strftime("%Y%m%d_%H%M%S")
        mode_label = "reasoning" if req.mode == "reasoning" else "simulation"
        lang_label = req.lang

        if req.format == "html":
            html_content = build_report_html(req)
            filename = f"quantum_report_{mode_label}_{lang_label}_{timestamp}.html"
            filepath = os.path.join(export_dir, filename)
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(html_content)
            return FileResponse(filepath, media_type="text/html", filename=filename)

        elif req.format == "docx":
            docx_bytes = build_report_docx(req)
            filename = f"quantum_report_{mode_label}_{lang_label}_{timestamp}.docx"
            filepath = os.path.join(export_dir, filename)
            with open(filepath, "wb") as f:
                f.write(docx_bytes)
            return FileResponse(filepath, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=filename)

        elif req.format == "pdf":
            # 使用weasyprint从HTML生成PDF
            filename = f"quantum_report_{mode_label}_{lang_label}_{timestamp}.pdf"
            filepath = os.path.join(export_dir, filename)
            try:
                from weasyprint import HTML as WeasyHTML
                html_content = build_report_html(req)
                WeasyHTML(string=html_content).write_pdf(filepath)
                return FileResponse(filepath, media_type="application/pdf", filename=filename)
            except Exception as pdf_error:
                print(f"[导出] WeasyPrint PDF生成失败，使用ReportLab兜底: {pdf_error}")
                pdf_bytes = build_report_pdf(req)
                with open(filepath, "wb") as f:
                    f.write(pdf_bytes)
                return FileResponse(
                    filepath,
                    media_type="application/pdf",
                    filename=filename,
                    headers={"X-Export-Fallback": "weasyprint-to-reportlab"}
                )

        else:
            raise HTTPException(status_code=400, detail=f"不支持的格式: {req.format}")

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"报告生成失败: {str(e)}")


@app.api_route("/{static_path:path}", methods=["GET", "HEAD"])
async def serve_static_file(static_path: str):
    """在同一个 6122 端口上服务前端页面和静态资源。"""
    file_path, media_type = resolve_static_file(static_path)
    return FileResponse(file_path, media_type=media_type)


# ============================================================
# 启动入口
# ============================================================

if __name__ == "__main__":
    import uvicorn
    print("[启动] 量子知识图谱推演引擎")
    print(f"[配置] 监听: {DEFAULT_HOST}:{DEFAULT_PORT}")
    print(f"[配置] 端口: {DEFAULT_PORT}")
    print(f"[配置] DeepSeek Model: {DEEPSEEK_MODEL}")
    print(f"[配置] Graph Data: {GRAPH_DATA_PATH}")
    uvicorn.run(app, host=DEFAULT_HOST, port=DEFAULT_PORT)
